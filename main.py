from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from typing import Any, Dict, List, Optional, Literal
from io import BytesIO
from copy import deepcopy
import json
import re
import time
import os
import secrets
import tempfile
import subprocess

FRONTEND_BASE_URL = os.getenv("FRONTEND_BASE_URL", "http://localhost:8000")

import smtplib
from email.message import EmailMessage
from email.utils import formataddr


from fastapi.responses import Response
from docx import Document

from urllib.parse import urlencode
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import datetime as dt
import jwt
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from fastapi import Depends, HTTPException, status, Security
from pydantic import BaseModel
from passlib.hash import pbkdf2_sha256
from fastapi.responses import RedirectResponse
import httpx

from database import SessionLocal, engine
from models import Base, User,PasswordResetToken


from model_engine import (
    load_document_from_bytes,
    parse_docx_to_model,
    apply_model_and_generate_docx,
    build_cover_letter_docx,
)

from llm_client import get_chat_completion

LLM_PROVIDER_PRIORITY = ["gemini", "perplexity", "openai"]



# ---------- AUTH Pydantic models ----------
class RegisterRequest(BaseModel):
    email: str
    password: str

class LoginRequest(BaseModel):
    email: str
    password: str

class AuthResponse(BaseModel):
    token: str
    email: str
    user_id: int

class ForgotPasswordRequest(BaseModel):
    email: str

class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str
    confirm_password: str

class DetectionIssue(BaseModel):
    id: str
    section: str             # "summary", "skills", "experience", "education", "certifications", "global"
    item_id: Optional[str]   # e.g. "summary_1", "exp_2_bullet_3"
    severity: Literal["info", "warning", "error"]
    code: str                # e.g. "WEAK_VERB", "BULLET_TOO_LONG"
    message: str
    suggestion: Optional[str] = None


class DetectionScores(BaseModel):
    overall_quality: int
    bullet_strength: int
    keyword_alignment: int
    clarity: int


class DetectionReport(BaseModel):
    template_case: Optional[str] = None
    scores: DetectionScores
    issues: List[DetectionIssue]


class DetectRequest(BaseModel):
    model_json: str
    job_description: Optional[str] = None



load_dotenv()

JWT_SECRET = os.getenv("JWT_SECRET")
JWT_ALGORITHM = os.getenv("JWT_ALGORITHM", "HS256")

if not JWT_SECRET:
    raise RuntimeError("JWT_SECRET environment variable is not set")

GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")
GOOGLE_REDIRECT_URI = os.getenv("GOOGLE_REDIRECT_URI")

# Create all tables
Base.metadata.create_all(bind=engine)


# FastAPI app is already created in your file, keep it as is:
# app = FastAPI()

# --- DB dependency ---
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# --- JWT helpers ---
def create_token(user_id: int) -> str:
    payload = {
        "sub": user_id,
        "exp": dt.datetime.utcnow() + dt.timedelta(days=7),
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)
    return token

security = HTTPBearer()

def get_current_user(
    credentials: HTTPAuthorizationCredentials = Security(security),
    db: Session = Depends(get_db),
):
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id: int = payload.get("sub")
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired token",
        )

    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="User not found",
        )
    return user




def _clean_slug(s: str, fallback: str) -> str:
    """
    Turn 'Abdul Sattar Shaik' -> 'AbdulSattarShaik'
    Remove spaces and weird characters so it's safe for filenames.
    """
    if not isinstance(s, str):
        return fallback
    s = s.strip()
    if not s:
        return fallback
    s = re.sub(r"[^A-Za-z0-9]+", "", s)
    return s or fallback


def build_filename_from_model(model: Dict[str, Any], kind: str = "resume") -> str:
    """
    Build a nice filename like:
      abdul_sattar_shaik_edi_specialist_resume.docx
      abdul_sattar_shaik_edi_specialist_cover_letter.docx
    based on model['header'].

    Supports both:
      - older header["name"] / header["title"] fields (if UI ever sets them)
      - current header["lines"] = [name, role/title, contact...]
    """
    header = model.get("header") or {}
    raw_name = ""
    raw_role = ""

    if isinstance(header, dict):
        # Preferred: explicit fields if present
        raw_name = (header.get("name") or "").strip()
        raw_role = (header.get("title") or header.get("role") or "").strip()

        # New layout from model_engine: header["lines"] = [name, role/title, contact...]
        lines = header.get("lines") or []
        if isinstance(lines, list):
            # Name from first non-empty line
            if not raw_name:
                for ln in lines:
                    s = str(ln).strip()
                    if s:
                        raw_name = s
                        break

            # Role/title from a non-contact-like line after the name
            if not raw_role:
                def _looks_contact_like(text: str) -> bool:
                    t = text.lower()
                    if "@" in t or "linkedin" in t or "http" in t or "www." in t:
                        return True
                    digits = sum(ch.isdigit() for ch in t)
                    if digits >= 3 and any(sep in t for sep in ("-", "(", ")", "+")):
                        return True
                    return False

                for ln in lines[1:]:
                    s = str(ln).strip()
                    if not s:
                        continue
                    if _looks_contact_like(s):
                        continue
                    raw_role = s
                    break

    if not raw_name:
        raw_name = "resume"

    def _to_pascal_case(s: str) -> str:
    # Remove non-alphanumeric characters and convert to PascalCase
         parts = re.split(r"[^A-Za-z0-9]+", s)
         return "".join(p.capitalize() for p in parts if p)


    name_slug = _to_pascal_case(raw_name)
    role_slug = _to_pascal_case(raw_role)

    base = name_slug
    if role_slug:
       base = f"{name_slug}_{role_slug}"
    else:
       base = name_slug

    if kind == "cover_letter":
     return f"{base}_CoverLetter.docx"
    else:
        return f"{base}_Resume.docx"

def parse_llm_providers_field(raw: Optional[str]) -> List[str]:
    """
    Parse a comma-separated 'llm_providers' form field from the frontend
    into a normalized list of provider identifiers: 'openai', 'gemini', 'perplexity'.
    """
    if not raw:
        return []

    parts = [p.strip().lower() for p in raw.split(",")]
    providers: List[str] = []

    for p in parts:
        if p in ("openai", "chatgpt", "gpt"):
            val = "openai"
        elif p in ("gemini", "google"):
            val = "gemini"
        elif p in ("perplexity", "pplx", "perplixity"):
            val = "perplexity"
        else:
            continue
        if val not in providers:
            providers.append(val)

    return providers


def choose_llm_from_list(candidates: List[str]) -> Dict[str, str]:
    """
    Given a list of provider candidates from the UI, choose a single
    provider + model to use for this request.

    Priority:
      - gemini
      - perplexity
      - openai

    If no candidates are provided, fall back to environment defaults.
    """
    if not candidates:
        provider = os.getenv("LLM_PROVIDER", "openai").lower()
        model = os.getenv("LLM_MODEL", "gpt-4.1-mini")
        return {"provider": provider, "model": model}

    provider = candidates[0]
    for pref in LLM_PROVIDER_PRIORITY:
        if pref in candidates:
            provider = pref
            break

    if provider == "openai":
        model = os.getenv("LLM_MODEL", "gpt-4.1-mini")
    elif provider == "gemini":
        # Match llm_client.py default: current Gemini text model
        model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
    elif provider == "perplexity":
        # Match llm_client.py default: official Sonar model
        model = os.getenv("PERPLEXITY_MODEL", "sonar")
    else:
        provider = "openai"
        model = os.getenv("LLM_MODEL", "gpt-4.1-mini")

    return {"provider": provider, "model": model}




app = FastAPI()

# Allow frontend (file:// or localhost) to call backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # OK for local dev
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)


# =======================
# SMTP EMAIL (PASSWORD RESET)
# =======================

SMTP_HOST = os.getenv("SMTP_HOST")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")
SMTP_FROM = os.getenv("SMTP_FROM") or SMTP_USER
SMTP_FROM_NAME = os.getenv("SMTP_FROM_NAME", "SATTAR JD + Resume")

def send_email_smtp(to_email: str, subject: str, body: str) -> None:
    """
    Simple SMTP email sender using environment variables.

    This is synchronous and blocking, which is fine for low traffic.
    """
    if not (SMTP_HOST and SMTP_PORT and SMTP_USER and SMTP_PASSWORD and SMTP_FROM):
        # You can raise or just log; for now we raise so you see the misconfig.
        raise RuntimeError("SMTP is not configured. Set SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD, SMTP_FROM.")

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = formataddr((SMTP_FROM_NAME, SMTP_FROM))
    msg["To"] = to_email
    msg.set_content(body)

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.send_message(msg)


def send_password_reset_email(to_email: str, reset_url: str) -> None:
    """
    Build a simple password-reset email and send it to the user.
    """
    subject = "Reset your password – SATTAR JD + Resume"

    body = f"""Hello,

We received a request to reset the password for your account.

To reset your password, click the link below or paste it into your browser:

{reset_url}

If you did not request a password reset, you can safely ignore this email.

Thanks,
SATTAR JD + Resume
"""

    send_email_smtp(to_email, subject, body)



# ============================================================
# =============== GPT REWRITE HELPERS (TEXT) =================
# ============================================================

def build_rewrite_payload(job_description: str, model: Dict[str, Any]) -> Dict[str, Any]:
    """
    Build a clean JSON payload for GPT from our internal model.

    Important:
    - Only include summary, skills, experience.
    - Add stable IDs so we can map changes back without reordering.
    - For experience, also include optional environment / role / project_description
      as {label, value} objects so GPT can update the tech stack & descriptions.
    """
    payload: Dict[str, Any] = {
        "job_description": job_description,
        "schema_version": 1,
    }

    # Summary bullets
    summary = model.get("summary")
    summary_payload: List[Dict[str, Any]] = []

    # Our normal schema: {"bullets": ["...", "..."]}
    if isinstance(summary, dict):
        bullets = summary.get("bullets") or []
        for i, bullet in enumerate(bullets, start=1):
            if not isinstance(bullet, str):
                continue
            text = bullet.strip()
            if not text:
                continue
            summary_payload.append({
                "id": i,
                "text": text,
            })

    # Legacy / safety: if summary is already a list, still support it
    elif isinstance(summary, list):
        for i, item in enumerate(summary, start=1):
            if isinstance(item, dict):
                text = (item.get("text") or "").strip()
                sid = item.get("id", i)
            else:
                text = str(item).strip()
                sid = i
            if not text:
                continue
            summary_payload.append({
                "id": sid,
                "text": text,
            })

    if summary_payload:
        payload["summary"] = summary_payload

    # Skills table
    skills = model.get("skills")
    if isinstance(skills, dict):
        rows = skills.get("rows") or []
        skills_rows: List[Dict[str, Any]] = []

        for i, row in enumerate(rows, start=1):
            label = None
            value = None

            # Normal shape from parser: ["Label", "Value"]
            if isinstance(row, list) and len(row) >= 2:
                label, value = row[0], row[1]

            # Legacy shape: {"label": "...", "value": "..."}
            elif isinstance(row, dict):
                label = row.get("label")
                value = row.get("value")

            if not isinstance(value, str):
                continue

            skills_rows.append({
                "id": i,
                "label": (label or "") if isinstance(label, str) else "",
                "value": value,
            })

        if skills_rows:
            payload["skills"] = {"rows": skills_rows}

    # Experience
    experience = model.get("experience")
    if isinstance(experience, list):
        payload["experience"] = []
        for j, job in enumerate(experience):
            if not isinstance(job, dict):
                continue
            bullets = job.get("bullets") or []

            exp_entry: Dict[str, Any] = {
                "id": j + 1,
                "header": job.get("header", ""),
                "bullets": [
                    {"id": i + 1, "text": b} for i, b in enumerate(bullets)
                ],
            }

            # Optional label/value fields
            env_obj = job.get("environment")
            if isinstance(env_obj, dict):
                exp_entry["environment"] = {
                    "label": env_obj.get("label", ""),
                    "value": env_obj.get("value", ""),
                }

            role_obj = job.get("role")
            if isinstance(role_obj, dict):
                exp_entry["role"] = {
                    "label": role_obj.get("label", ""),
                    "value": role_obj.get("value", ""),
                }

            proj_obj = job.get("project_description")
            if isinstance(proj_obj, dict):
                exp_entry["project_description"] = {
                    "label": proj_obj.get("label", ""),
                    "value": proj_obj.get("value", ""),
                }

            payload["experience"].append(exp_entry)

    return payload


def collect_forbidden_terms_from_model(job_description: str, model: Dict[str, Any]) -> List[str]:
    """
    Build a list of 'old domain' tokens that appear in the resume model
    but do NOT appear in the JD. These will be passed to GPT as
    'forbidden_terms' so it avoids carrying over EDI/network/etc
    when the new JD is Java/Data/etc.
    """
    jd_low = job_description.lower()

    def tokenize(text: str) -> List[str]:
        tokens = re.split(r"[^A-Za-z0-9\+\#\.]+", text)
        tokens = [t for t in tokens if t]
        return tokens

    resume_tokens: List[str] = []

    # Summary
    summary = model.get("summary") or {}
    if isinstance(summary, dict):
        bullets = summary.get("bullets") or []
        for b in bullets:
            if isinstance(b, str):
                t = b.strip()
                if t:
                    resume_tokens.extend(tokenize(t))
    elif isinstance(summary, list):
        # Legacy support
        for s in summary:
            if isinstance(s, dict):
                t = (s.get("text") or "").strip()
                if t:
                    resume_tokens.extend(tokenize(t))

    # Skills
    skills = model.get("skills") or {}
    rows = skills.get("rows") or []
    if isinstance(rows, list):
        for row in rows:
            label = None
            value = None

            # Normal shape: ["Label", "Value"]
            if isinstance(row, list) and len(row) >= 2:
                label, value = row[0], row[1]
            # Legacy shape: {"label": "...", "value": "..."}
            elif isinstance(row, dict):
                label = row.get("label")
                value = row.get("value")

            if isinstance(label, str):
                lab = label.strip()
                if lab:
                    resume_tokens.extend(tokenize(lab))
            if isinstance(value, str):
                val = value.strip()
                if val:
                    resume_tokens.extend(tokenize(val))

    # Experience
    for job in model.get("experience", []):
        if not isinstance(job, dict):
            continue
        header = (job.get("header") or "").strip()
        if header:
            resume_tokens.extend(tokenize(header))

        env = job.get("environment") or {}
        if isinstance(env, dict):
            val = (env.get("value") or "").strip()
            if val:
                resume_tokens.extend(tokenize(val))

        role = job.get("role") or {}
        if isinstance(role, dict):
            val = (role.get("value") or "").strip()
            if val:
                resume_tokens.extend(tokenize(val))

        proj = job.get("project_description") or {}
        if isinstance(proj, dict):
            val = (proj.get("value") or "").strip()
            if val:
                resume_tokens.extend(tokenize(val))

        bullets = job.get("bullets") or []
        for b in bullets:
            if isinstance(b, str):
                resume_tokens.extend(tokenize(b))

    # Now see which tokens are NOT in the JD at all.
    forbidden: List[str] = []
    seen = set()

    for tok in resume_tokens:
        low = tok.lower()
        if len(low) <= 1:
            continue
        if low in seen:
            continue
        if low not in jd_low:
            forbidden.append(tok)
            seen.add(low)

    return forbidden


def call_gpt_rewrite(
    payload: Dict[str, Any],
    llm_provider: str = "openai",
    llm_model: Optional[str] = None,
) -> Dict[str, Any]:

    """
    Call LLM to rewrite the resume sections for the JD.

    Behaviour:
    - Uses 'ats_keywords' if provided (or auto-extracts if missing).
    - Uses 'forbidden_terms' (old-domain words not in JD) to ban EDI/network/etc
      when JD is Java/Data/etc.
    - Higher temperature (0.8) for more variation between runs (for non-5.1 models).
    - If payload['force_stronger'] is True, explicitly asks for a more aggressive rewrite.
    """
    jd = payload["job_description"]

    ats_keywords = payload.get("ats_keywords")
    force_all_keywords = bool(payload.get("ats_force_all"))
    forbidden_terms = payload.get("forbidden_terms") or []
    disable_ats = bool(payload.get("disable_ats"))
    force_stronger = bool(payload.get("force_stronger"))
    target_role_hint = payload.get("target_role_hint")  # <-- new hint (e.g. "Data Engineer")

    summary = payload.get("summary")
    skills = payload.get("skills")
    experience = payload.get("experience")

    sections = []
    if isinstance(summary, list):
        sections.append("summary")
    if isinstance(skills, dict):
        sections.append("skills")
    if isinstance(experience, list):
        sections.append("experience")

    # Give GPT some notion of variation for multiple runs on same JD + resume
    variation_hint = f"run_{int(time.time() * 1000)}"

    # ---- ATS keyword extraction if needed ----
    if disable_ats:
        ats_keywords = []
    else:
        if not ats_keywords:
            # Auto-extract ATS keywords from JD if caller didn't pass any.
            from_text = f"JOB DESCRIPTION:\n{jd}"
            try:
                kw_resp = get_chat_completion(
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are an assistant that extracts important ATS keywords "
                                "from job descriptions."
                            ),
                        },
                        {
                            "role": "user",
                            "content": (
                                "Extract the TOP 40–60 most important technical and domain "
                                "keywords/key phrases from this job description. "
                                "Return them as a JSON object with a single key 'keywords', "
                                "whose value is an array of strings.\n\n"
                                f"{from_text}"
                            ),
                        },
                    ],
                    model=llm_model,
                    temperature=0.0,
                    json_mode=True,
                    provider=llm_provider,
                )
                data = json.loads(kw_resp)
                kw = data.get("keywords", [])
                cleaned_kw: List[str] = []
                for item in kw:
                    if isinstance(item, str):
                        s = item.strip()
                        if s:
                            cleaned_kw.append(s)
                ats_keywords = cleaned_kw
            except Exception:
                ats_keywords = []

    # ATS block
    ats_block = ""
    if ats_keywords:
        deduped = []
        seen_kw = set()
        for kw in ats_keywords:
            s = str(kw).strip()
            l = s.lower()
            if s and l not in seen_kw:
                seen_kw.add(l)
                deduped.append(s)

        ats_keywords = deduped

        ats_block = (
            "ATS KEYWORDS:\n"
            "- The following list are ATS keywords/keyphrases that should be well-covered in the rewritten resume.\n"
            "- When rewriting, prefer to integrate them NATURALLY into the summary, skills, and experience bullets.\n"
            "- Only if you absolutely cannot fit them naturally, you may add them to skills or environment.\n"
        )
        if force_all_keywords:
            ats_block += (
                "- FORCE MODE: You MUST ensure that EVERY keyword in the list appears at least once somewhere "
                "in the rewritten resume (summary, skills, experience bullets, or environment/role fields).\n"
            )
        else:
            ats_block += (
                "- NORMAL MODE: Try to cover MOST of these keywords, but do NOT spam them unnaturally.\n"
            )
        ats_block += "\nKEYWORDS:\n"
        for kw in ats_keywords:
            ats_block += f"- {kw}\n"
    else:
        ats_block = (
            "No explicit ATS keyword list was provided. You should still optimize the resume for ATS by leaning "
            "on critical technologies, tools, and domain language from the job description.\n"
        )

    # Forbidden terms block
    forbidden_block = ""
    if forbidden_terms:
        deduped_f = []
        seen_f = set()
        for t in forbidden_terms:
            s = str(t).strip()
            l = s.lower()
            if s and l not in seen_f:
                seen_f.add(l)
                deduped_f.append(s)
        forbidden_terms = deduped_f

        forbidden_block = (
            "FORBIDDEN TERMS (OLD DOMAIN / MUST AVOID):\n"
            "- The following tokens come from the the existing resume but DO NOT appear in the job description.\n"
            "- Treat these as 'old domain' or 'incorrect domain' terms that must be removed or replaced.\n"
            "- You MUST NOT use them in the rewritten resume unless the job description explicitly contains them.\n"
            "- If you need to express the concept, use the equivalent from the job description instead.\n\n"
        )
        for t in forbidden_terms:
            forbidden_block += f"- {t}\n"
    else:
        forbidden_block = (
            "There is no explicit forbidden-term list. However, you should STILL avoid keeping domain-specific "
            "terms that clearly do not belong to the new JD.\n"
        )

    force_stronger_block = ""
    if force_stronger:
        force_stronger_block = (
            "FORCE STRONGER REWRITE MODE:\n"
            "- This is a SECOND attempt because your previous rewrite was too similar to the original.\n"
            "- You MUST significantly change the wording and structure of each bullet compared to the input.\n"
            "- Do NOT reuse the same phrases or sentence structures from the original text.\n"
            "- Keep facts (companies, dates, locations, seniority, and overall years of experience) the same, but\n"
            "  aggressively rephrase how responsibilities and achievements are described.\n\n"
        )

    special_role_block = ""
    if target_role_hint:
        special_role_block = (
            "SPECIAL ROLE DOMAIN INSTRUCTION:\n"
            f"- The JD is for a {target_role_hint} style role.\n"
            f"- For every experience job, ensure the job header's role title and any 'role.value'\n"
            f"  field use a {target_role_hint} style title (e.g. 'Data Engineer', 'Senior Data Engineer'),\n"
            "  not old-domain titles like 'EDI Specialist', '.NET Developer', or 'Network Engineer'.\n"
            "- Also ensure each job's environment.value lists realistic technologies for this role, aligned\n"
            "  with the JD, instead of old-domain stacks or EDI/.NET/network-specific tools.\n\n"
        )

    # 🔑 ALWAYS build user_content here
    user_content = (
        "You are rewriting a resume to match a given Job Description (JD).\n\n"
        "There is an existing resume represented as JSON with sections:\n"
        "- summary: list of bullets (each has id, text)\n"
        "- skills: table with rows (id, label, value)\n"
        "- experience: list of jobs, where each job has:\n"
        "    id, header (company/location/dates/role), bullets, and optional\n"
        "    environment / role / project_description objects with {label, value}.\n\n"
        "YOUR GOAL:\n"
        "- Rewrite ALL the resume content so it is strongly aligned to the JD domain (e.g. Java, .NET, Data, EDI, etc).\n"
        "- Keep the structure and formatting stable so the DOCX template remains intact.\n"
        "- Preserve true facts like companies, dates, locations, and overall seniority/years of experience.\n"
        "- Remove or replace domain-specific terms from the old resume that conflict with the JD domain.\n\n"
        "HARD CONSTRAINTS (STRUCTURE):\n"
        "1) Do NOT change the number or order of summary bullets, skill rows, or experience jobs.\n"
        "   - Keep the same number and order of experience entries and their bullets.\n"
        "2) Do NOT add or remove experiences, companies, locations, degrees, or dates.\n"
        "3) You may ONLY change the text fields:\n"
        "   - summary[i].text\n"
        "   - skills.rows[i].label\n"
        "   - skills.rows[i].value\n"
        "   - experience[j].header\n"
        "   - experience[j].bullets[k].text\n"
        "   - experience[j].environment.value (if present)\n"
        "   - experience[j].role.value (if present)\n"
        "   - experience[j].project_description.value (if present)\n"
        "4) Never change or invent IDs. All 'id' values must remain exactly the same.\n"
        "5) YEARS OF EXPERIENCE RULE: If a bullet or header already mentions a number of years (e.g. '6+ years'),\n"
        "   KEEP that number. Do NOT reduce it to match the JD (e.g. do NOT change 6+ to 3+ just because the JD\n"
        "   says '3+ years'). You may rephrase around it, but do not downgrade the candidate.\n\n"
        "WRITING STYLE & ANTI-AI RULES (VERY IMPORTANT):\n"
        "1) Write in concise, professional, HUMAN language – like a real engineer describing their own work to another\n"
        "   engineer. The tone should feel practical and grounded, not like marketing or an AI template.\n"
        "2) Do NOT use chatty or AI-ish phrases such as: 'In this role', 'In my role', 'In this position',\n"
        "   'As a [title]', 'We', 'I', 'My responsibilities included', 'Duties included'.\n"
        "3) Start bullets with concrete actions, but avoid over-formal or template-like verbs.\n"
        "   - Prefer verbs like: Designed, Implemented, Developed, Fixed, Debugged, Tuned, Deployed,\n"
        "     Refactored, Improved, Automated, Migrated, Configured, Tested.\n"
        "   - Do NOT start bullets with: Led, Built, Adopted, Adept, Orchestrated/Orchestrates,\n"
        "     Facilitated, Authored, Aligned, Reconciled, 'Hands-on exposure to', or similar AI-ish openings.\n"
        "4) Do NOT start any bullet or sentence with:\n"
        "   - a comma ','\n"
        "   - 'In', 'On', 'This', 'There is', 'There are', 'Here is', or similar filler.\n"
        "   Always rewrite so bullets start directly with a clear, concrete action.\n"
        "5) Avoid over-used fluff or corporate buzzwords. The following phrases MUST NOT appear:\n"
        "   - 'results-driven', 'detail-oriented', 'highly motivated', 'dynamic professional', 'seasoned',\n"
        "   - 'passionate problem solver', 'overall', 'leveraged synergies', 'team player', 'proven track record',\n"
        "   - 'collaborated cross-functionally', 'worked on various projects', 'clear communicator',\n"
        "   - 'trained new team members', 'trained', 'mentored', 'authored' (as a generic filler),\n"
        "   - marketing-style lines like 'reliable professional', 'organized and effective under tight deadlines',\n"
        "     'thrives in high-pressure environments', or any similar AI-ish self-promotion.\n"
        "6) Avoid filler adverbs like 'successfully', 'efficiently', 'effectively' unless they add real meaning.\n"
        "7) Bullets must sound like real project work, not like a cover letter or marketing text.\n"
        "8) PUNCTUATION RULE: Do NOT use semicolons (';'), question marks ('?'), or exclamation marks ('!')\n"
        "   in bullets or skills values. Use simple sentences with periods and commas instead.\n\n"
        "CONTENT RULES (SUMMARY, SKILLS, EXPERIENCE):\n"
        "0) MINIMUM QUALIFICATIONS / LOGISTICS FROM JD:\n"
        "   - Do NOT turn logistical requirements from the JD into resume bullets or summary lines.\n"
        "   - Examples: onsite/remote/hybrid rules, days per week in office, specific city for the office,\n"
        "     background checks, work authorization statements, equal opportunity/legal statements,\n"
        "     benefits, or generic HR policy text.\n"
        "   - You may respect these constraints implicitly (e.g., not contradicting them), but you MUST NOT\n"
        "     create marketing-style bullets such as 'Reliable professional based in Manhattan, organized and\n"
        "     effective under tight deadlines and high-pressure situations'. Simply ignore those JD lines when\n"
        "     rewriting the resume content.\n\n"
        "1) SUMMARY:\n"
        "   - Rewrite EVERY summary bullet so it directly reflects the JD domain and responsibilities.\n"
        "   - Remove obviously irrelevant old-domain technologies unless they also appear in the JD.\n"
        "   - Aim for concise, high-impact bullets that sound human and professional.\n"
        "   - Avoid generic fluff like 'responsible for', 'worked on' unless needed for clarity.\n"
        "   - If the original summary has many bullets, you may compress similar ideas into stronger bullets,\n"
        "     but you MUST still return the same count of bullets (just with richer content).\n\n"
        "2) SKILLS:\n"
"   - The skills section is a 2-column TECH STACK table.\n"
"   - You are NOT creating a new generic skills section. You are UPDATING an existing, detailed skills table.\n"
"   - Keep the original row labels/headings EXACTLY the same (e.g. 'Programming Languages', '.NET & Web Technologies',\n"
"     'Banking & Payments Domain', 'XML & Data Transformation', 'EDI Transactions', 'Order Transactions').\n"
"     Do NOT invent new headings like 'Core Data Engineering Technologies' or 'Data Modernization & AI Techniques'.\n"
"   - For each row, UPDATE ONLY the value text (right side) so that it aligns with the JD domain, but:\n"
"       * Preserve the overall density and seniority. If the original row lists many tools/versions, your\n"
"         rewritten row should also be rich and detailed – NOT a short generic list.\n"
"       * Prefer to REPLACE old-domain tools with JD-appropriate tools, instead of collapsing the row into\n"
"         a tiny list. Example: swap out 'EDI/X12' for 'REST APIs / Spring Boot' rather than shrinking the row.\n"
"   - The RIGHT side of each row must be a comma-separated list of tools, technologies, standards, codes, or\n"
"     short noun phrases (1–4 words each). It must NOT contain full sentences.\n"
"   - Do NOT use verbs in skills cells (no 'managing', 'ensuring', 'driving', 'performing', 'handling',\n"
"     'supporting', 'minimizing financial loss', 'improving processes', etc.). Responsibilities and outcomes\n"
"     belong in the EXPERIENCE section, not in SKILLS.\n"
"   - Never leave rows empty; if a row's original value is irrelevant, replace it with JD-appropriate skills\n"
"     of similar depth (comparable number of tools/frameworks/versions).\n"
"   - Do NOT add comments like 'Removed to align with the JD'; skill rows must contain only real skills.\n"
"   - SKILL ROW DUPLICATION RULE: avoid creating rows that have EXACTLY the same label + value as another row.\n"
"     If a row would be duplicate, change the value so each row feels distinct (different group of tools).\n"
"   - Do NOT turn skills into vague phrases like 'Data Engineering Automation' or 'Collaboration & Documentation'\n"
"     unless the original row was like that. Focus on concrete tools, technologies, standards, platforms, and codes.\n"
"   - SPECIAL RULE FOR TRANSACTION / PROTOCOL ROWS:\n"
"       * If a row label contains words like 'EDI Transactions', 'Order Transactions', or 'Routing Protocols',\n"
"         you MUST preserve the numeric codes / protocol names (e.g. '837I/837P/837D, 835, 270/271, 276/277, 999, 277CA, TA1'\n"
"         or '850 (PO), 855 (PO Ack), 810 (Invoice), 856 (ASN), 846 (Inventory)', or 'OSPF, EIGRP, BGP').\n"
"       * You may ADD short names in parentheses (e.g. '850 (PO)', '855 (PO Ack)'), but you must NOT remove or\n"
"         replace the codes themselves.\n"
"   - STYLE EXAMPLES (you must follow this style):\n"
"       * EDI Standards      ANSI X12, EDIFACT, HIPAA\n"
"       * EDI Transactions   837I/837P/837D, 835, 270/271, 276/277, 999, 277CA, TA1\n"
"       * Order Transactions 850 (PO), 855 (PO Ack), 810 (Invoice), 856 (ASN), 846 (Inventory)\n"
"       * Programming Languages  C, C++, Java, C#, Python\n""\n"
 "3) EXPERIENCE:\n"
        "   - For EACH job:\n"
        "     a) HEADER:\n"
        "        - Preserve company name, location, and dates exactly.\n"
        "        - If the header includes a role title (e.g. 'EDI Developer'), update it to a JD-aligned role\n"
        "          (e.g. 'Data Engineer', '.NET Developer'), while keeping seniority accurate.\n"
        "        - For a given JD, keep the role domain CONSISTENT across all jobs.\n"
        "          For example, if the JD is for a Data Engineer, every job header should use a Data Engineer\n"
        "          style title (e.g. 'Data Engineer', 'Senior Data Engineer'), and MUST NOT switch to\n"
        "          'Java Developer', 'Backend Developer', etc.\n"
        "     b) BULLETS:\n"
        "        - Rewrite EVERY bullet to emphasize responsibilities and achievements that match the JD.\n"
        "        - Remove or replace old-domain technical terms with equivalent JD-domain technologies.\n"
        "        - Keep achievements realistic and consistent with the role and industry.\n"
        "        - Do NOT simply append 1–2 JD buzzwords; fully rewrite the sentence into the new domain.\n"
        "        - Use a natural, human tone, as if the candidate really did work in that JD domain.\n"
        "        - Avoid robotic or repetitive phrasing; vary sentence structure so bullets read like\n"
        "          natural notes from a real engineer, not AI-generated templates.\n"
        "     c) ENVIRONMENT / ROLE / PROJECT DESCRIPTION:\n"
        "        - For each job, if environment/role/project_description objects exist:\n"
        "          * environment.value: update this to list the key tools/technologies for that job in the JD domain.\n"
        "            Environment values across different jobs MUST NOT be copy-paste identical.\n"
        "            It is fine if they share 40–60% of technologies, but each job's environment should have at least\n"
        "            2–3 unique tools/platforms/frameworks that make sense for that job.\n"
        "          * role.value: align to the JD role (e.g. 'Senior Java Developer', 'Data Analyst') as appropriate.\n"
        "          * project_description.value: rewrite to describe the project in terms that match the JD.\n\n"
        "BOLDING RULES:\n"
        "1) For EVERY summary bullet and EVERY experience bullet in EVERY job, include multiple short bold\n"
        "   technical phrases where it makes sense (typically 2–3 bold chunks per bullet).\n"
        "2) Treat the ATS KEYWORDS list (if provided above) as your PRIMARY source of bold phrases:\n"
        "   - Whenever you naturally use one of those keywords (especially tools, frameworks, platforms,\n"
        "     databases, cloud services), wrap it in **bold** the FIRST time it appears in the resume.\n"
        "   - Aim for at least 70–80% of the ATS keywords to appear in bold at least once somewhere across\n"
        "     summary, skills, and experience bullets.\n"
        "3) Focus bolding on concrete technologies and domain terms, not generic verbs or fluff.\n"
        "   Examples: '**Python, PySpark**', '**Kafka streams**', '**SQL Server 2022**', '**AWS S3, Glue, EMR**',\n"
        "   '**X12 837/835**', '**Snowflake warehouses**'.\n"
        "4) Do NOT bold entire sentences or whole bullets. Bold should help a reviewer scan for key tech quickly.\n"
        "   As a rough guide, avoid bolding more than ~30–40% of the characters in a single bullet.\n"
        "5) Spread bolding evenly so that later bullets and later jobs also contain bold ATS keywords,\n"
        "   not just the first job or first few lines.\n\n"
        "VARIATION / MULTIPLE RUNS:\n"
        "1) You may be called multiple times with the SAME JD and SAME resume.\n"
        "2) Each time, you MUST produce slightly different wording and examples (while staying truthful and aligned).\n"
        "3) Maintain the same structure and facts, but vary phrasing and which details you emphasize.\n\n"
        "JSON OUTPUT FORMAT (VERY IMPORTANT):\n"
        "1) You MUST output a single JSON object with ONLY these keys:\n"
        "      'schema_version' (optional, copy from input if present)\n"
        "      'summary' (if present in input)\n"
        "      'skills'  (if present in input)\n"
        "      'experience' (if present in input)\n"
        "2) For each section, preserve the 'id' values and overall structure exactly.\n"
        "3) Do NOT wrap the JSON in markdown. No backticks. No extra commentary.\n\n"
        f"{ats_block}\n"
        f"{forbidden_block}\n"
        f"{force_stronger_block}"
        f"{special_role_block}"
        "Now here is the Job Description (JD):\n\n"
        f"{jd}\n\n"
        "And here is the existing resume JSON you must rewrite:\n"
        f"{json.dumps(payload, indent=2)}\n\n"
        f"Variation hint for this run: {variation_hint}\n"
        "Return ONLY the rewritten JSON object for summary, skills, and experience.\n"
    )

    response = get_chat_completion(
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a senior technical resume writer and ATS optimization expert. "
                    "You ALWAYS return valid JSON when asked for JSON, with no surrounding markdown."
                ),
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
        model=llm_model,
        temperature=1.1,
        json_mode=True,
        provider=llm_provider,
    )

    try:
        data = json.loads(response)
        if not isinstance(data, dict):
            raise ValueError("LLM returned JSON that is not an object.")
    except json.JSONDecodeError as e:
        raise ValueError(f"LLM rewrite did not return valid JSON: {e}") from e

    return data

def generate_cover_letter_text(
    model: Dict[str, Any],
    job_description: str,
    llm_provider: str,
    llm_model: str,
) -> str:
    """
    Use the current structured resume model + JD to generate a tailored cover letter.

    - Uses the updated resume model (after any rewrites).
    - Does NOT invent new companies, roles, or dates.
    - Avoids AI/meta language and generic buzzword fluff.
    """
    # Flatten model into readable resume text for the prompt
    try:
        resume_text = model_to_plain_text_for_ats(model)
    except Exception:
        # Fallback: raw JSON if flattening fails for some reason
        resume_text = json.dumps(model, indent=2)

    system_msg = (
        "You are an experienced hiring manager writing concise, professional, "
        "ATS-aware cover letters for job applications. "
        "You receive a candidate's resume and a target job description. "
        "Write a polished, natural-sounding cover letter that:\n"
        "- Stays truthful to the resume (no invented companies, roles, dates, or tools).\n"
        "- Focuses on the most relevant experience and skills for this specific job.\n"
        "- Uses clear, direct language (avoid generic buzzword fluff).\n"
        "- Does NOT mention AI, language models, or how the letter was generated.\n"
        "- Does NOT include placeholder fields like [Date], [Company Name], [Hiring Manager], etc.\n"
        "- Uses American English and a professional but warm tone.\n"
        "- Fits comfortably on one page.\n"
        "Formatting rules:\n"
        "- Do not add a placeholder date line. Omit the date completely.\n"
        "- Do not add a separate line that just says 'Hiring Manager' before the greeting.\n"
        "- Start the letter body with the greeting, e.g. 'Dear Hiring Manager,' on its own line.\n"
        "- End with a standard closing like 'Sincerely,' followed by the candidate's name on the next line."
    )

    user_msg = (
        "JOB DESCRIPTION:\n"
        f"{job_description.strip()}\n\n"
        "CANDIDATE RESUME (FLATTENED TEXT):\n"
        f"{resume_text}\n\n"
        "Using ONLY information that appears in the resume above, "
        "write a tailored cover letter for this role. "
        "Address it generically to 'Hiring Manager' (or '<Company> Hiring Manager' "
        "if the company name is clearly present in the job description). "
        "Do NOT include any placeholder fields like [Date] or [Company Name]. "
        "Do NOT add a separate address block; just write the greeting and paragraphs.\n\n"
        "Use 3–5 short paragraphs: a clear opening, one or two body paragraphs tying "
        "experience to the JD, and a concise closing paragraph."
    )

    raw = get_chat_completion(
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        provider=llm_provider,
        model=llm_model,
        temperature=0.4,
        json_mode=False,
    )

    text = raw.strip()

    # Backup cleanup: remove any leftover [Date]-style placeholder lines
    cleaned_lines: List[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        # Kill [Date] / [date] / [DATE] etc.
        if stripped.lower() in ("[date]", "[today's date]", "[today’s date]"):
            continue
        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


def is_protected_bullet(text: str) -> bool:
    """
    Return True if this bullet is effectively a heading/meta line that we never want
    GPT or post-processing to override (e.g. 'Responsibilities:', 'Environment:', 'WORK HISTORY').

    We are intentionally a bit generous here so we don't accidentally lose important headings.
    """
    if not text:
        return False

    t = text.strip()
    if not t:
        return False

    upper = t.upper()
    lower = t.lower()

    # Exact known headings (case-insensitive)
    HARD_HEADINGS = {
        "RESPONSIBILITIES:",
        "RESPONSIBILITIES",
        "ENVIRONMENT:",
        "ENVIRONMENT",
        "ROLE:",
        "ROLE",
        "PROJECT DESCRIPTION:",
        "PROJECT DESCRIPTION",
        "DESCRIPTION:",
        "DESCRIPTION",
        "WORK HISTORY",
        "WORK EXPERIENCE",
        "EXPERIENCE",
        "TECHNICAL SKILLS",
        "SKILLS",
        "EDUCATION",
        "EDUCATION DETAILS",
        "CERTIFICATIONS",
        "PROFESSIONAL SUMMARY",
        "SUMMARY",
    }
    if upper in HARD_HEADINGS:
        return True

    # Short heading ending with ':'  → treat as protected (e.g. 'Key Responsibilities:')
    if t.endswith(":") and len(t.split()) <= 4:
        return True

    # Heuristic: if it's a short line and contains these words, treat as heading
    heading_keywords = [
        "responsibilit",      # responsibilities / responsibility
        "environment",
        "work history",
        "work experience",
        "education",
        "technical skills",
        "project description",
        "summary",
    ]
    if len(t.split()) <= 5:
        for kw in heading_keywords:
            if kw in lower:
                return True

    return False


def _merge_header_preserving_structure(
    original_header: str,
    rewritten_header: str,
    has_explicit_role_line: bool,
) -> str:
    """
    Safely merge a rewritten header into the original:

    - If the job already has a separate Role: line, we DO NOT touch the header at all.

    - If there is NO explicit Role: line:
        * We treat the text before the first comma in the ORIGINAL header as the role/title.
        * We take the text before the first comma in the REWRITTEN header as the new role/title.
        * We replace ONLY that prefix, and keep the rest (company, location, dates) exactly as in the original.
    """
    if not isinstance(original_header, str):
        return rewritten_header if isinstance(rewritten_header, str) else original_header
    if not isinstance(rewritten_header, str):
        return original_header

    original_header = original_header.strip()
    rewritten_header = rewritten_header.strip()

    # If there is already an explicit Role: line, do NOT change header at all.
    if has_explicit_role_line:
        return original_header

    # Try to swap only the role/title before the first comma.
    idx_comma = original_header.find(",")
    if idx_comma == -1:
        # No comma, too risky to do surgery → trust rewritten header.
        return rewritten_header

    orig_role_part = original_header[:idx_comma].strip()
    rest = original_header[idx_comma:]  # includes comma + company/location/dates

    # New role candidate from rewritten header
    new_role_candidate = rewritten_header.split(",", 1)[0].strip()
    if not new_role_candidate:
        # Nothing usable → keep original header
        return original_header

    # If GPT didn't actually change the role text, just keep original header
    if new_role_candidate == orig_role_part:
        return original_header

    # Build new header: new role + original rest (company/location/dates untouched)
    return f"{new_role_candidate}{rest}"


def merge_rewrite_into_model(original: Dict[str, Any], rewrite: Dict[str, Any]) -> Dict[str, Any]:
    """
    Merge GPT rewrite JSON back into our original model, using IDs.

    Rules:
    - If GPT skips a section or item, we keep the original text.
    - If GPT adds weird IDs, we ignore them.
    - We keep the same number and order of bullets/rows/jobs as in the original.
    - Environment lines, role, project_description are merged via label/value,
      but the label text itself is kept from the original.
    - EXTRA: bullets that are sub-headings like 'Responsibilities:', 'Role:', 'Project Description:',
      'Description:', 'EDUCATION DETAILS', 'TECHNICAL SKILLS', 'WORK HISTORY', etc. are NEVER changed.
    - NEW: headers are now merged in a safer way using _merge_header_preserving_structure.
    """
    model = deepcopy(original)

    def merge_label_value_field(field_name: str) -> None:
        """
        For environment / role / project_description:
        - Keep the original label (e.g. 'Environment:')
        - If GPT provides a non-empty value, use it.
        - Otherwise keep the original value.
        """
        orig_field = orig_job.get(field_name)
        r_field = rj.get(field_name)

        if not isinstance(orig_field, dict):
            return

        base_label = orig_field.get("label", "")
        base_value = orig_field.get("value", "")

        new_value = None
        if isinstance(r_field, dict):
            candidate = r_field.get("value")
            if isinstance(candidate, str):
                candidate = candidate.strip()
                if candidate:
                    new_value = candidate

        # Fall back to original if GPT didn't supply anything useful
        if new_value is not None:
            job_new[field_name] = {"label": base_label, "value": new_value}
        else:
            job_new[field_name] = {"label": base_label, "value": base_value}

    # ---- MERGE SUMMARY ----
    summary_rewrite = rewrite.get("summary")
    if isinstance(summary_rewrite, list):
        base_summary = model.get("summary")

        if isinstance(base_summary, dict):
            old_bullets = base_summary.get("bullets") or []
        elif isinstance(base_summary, list):
            old_bullets = []
            for item in base_summary:
                if isinstance(item, dict):
                    txt = (item.get("text") or "").strip()
                    old_bullets.append(txt)
        else:
            old_bullets = []

        r_map: Dict[int, str] = {}
        for item in summary_rewrite:
            if not isinstance(item, dict):
                continue
            try:
                sid = int(item.get("id"))
            except (TypeError, ValueError):
                continue
            text_val = item.get("text")
            if isinstance(text_val, str):
                text_val = text_val.strip()
                if text_val:
                    r_map[sid] = text_val

        new_bullets: List[str] = []

        if old_bullets:
            for i, old in enumerate(old_bullets, start=1):
                new_bullets.append(r_map.get(i, old))

            max_existing = len(old_bullets)
            for sid in sorted(r_map.keys()):
                if sid > max_existing:
                    new_bullets.append(r_map[sid])
        else:
            for sid in sorted(r_map.keys()):
                new_bullets.append(r_map[sid])

        model["summary"] = {"bullets": new_bullets}

    # ---- MERGE SKILLS ----
    if "skills" in model and isinstance(model["skills"], dict):
        orig_skills = model["skills"]
        rewrite_skills = rewrite.get("skills")
        if isinstance(rewrite_skills, dict):
            orig_rows = orig_skills.get("rows") or []
            r_rows = rewrite_skills.get("rows") or []

            r_map: Dict[int, Dict[str, Any]] = {}
            for row in r_rows:
                if not isinstance(row, dict):
                    continue
                try:
                    rid = int(row.get("id"))
                except (TypeError, ValueError):
                    continue
                r_map[rid] = row

            new_rows = []
            for i, row in enumerate(orig_rows, start=1):
                r_row = r_map.get(i)

                if isinstance(row, list):
                    label = row[0] if len(row) >= 1 else ""
                    value = row[1] if len(row) >= 2 else ""

                    new_value = value
                    if isinstance(r_row, dict):
                        candidate = r_row.get("value", value)
                        if isinstance(candidate, str):
                            if "Removed to align" not in candidate:
                                new_value = candidate

                    new_rows.append([label, new_value])

                elif isinstance(row, dict):
                    rid = row.get("id") or i
                    label = row.get("label") or ""
                    value = row.get("value") or ""
                    new_value = value
                    if isinstance(r_row, dict):
                        candidate = r_row.get("value", value)
                        if isinstance(candidate, str) and "Removed to align" not in candidate:
                            new_value = candidate
                    new_rows.append(
                        {
                            "id": rid,
                            "label": label,
                            "value": new_value,
                        }
                    )
                else:
                    new_rows.append(row)

            orig_skills["rows"] = new_rows
            model["skills"] = orig_skills

    # ---- MERGE EXPERIENCE ----
    if "experience" in model and isinstance(model["experience"], list):
        orig_exp = model["experience"]
        r_exp = rewrite.get("experience")
        if isinstance(r_exp, list):
            r_map = {}
            for item in r_exp:
                if not isinstance(item, dict):
                    continue
                try:
                    jid = int(item.get("id"))
                except (TypeError, ValueError):
                    continue
                r_map[jid] = item

            new_exp = []
            for j, orig_job in enumerate(orig_exp):
                if not isinstance(orig_job, dict):
                    new_exp.append(orig_job)
                    continue
                jid = orig_job.get("id") or (j + 1)
                job_new = deepcopy(orig_job)

                if isinstance(jid, int) and jid in r_map:
                    rj = r_map[jid]

                    # --- HEADER MERGE (fixed) ---
                    r_header = rj.get("header")
                    if isinstance(r_header, str):
                        orig_header = orig_job.get("header", "")
                        has_explicit_role = isinstance(orig_job.get("role"), dict)
                        job_new["header"] = _merge_header_preserving_structure(
                            orig_header,
                            r_header,
                            has_explicit_role,
                        )

                    # --- ENV / ROLE / PROJECT_DESCRIPTION ---
                    for fld in ("environment", "role", "project_description"):
                        merge_label_value_field(fld)

                    # --- BULLETS ---
                    orig_bullets = orig_job.get("bullets") or []
                    b_map = {i + 1: text for i, text in enumerate(orig_bullets)}

                    for b in rj.get("bullets", []):
                        if not isinstance(b, dict):
                            continue
                        try:
                            bid = int(b.get("id"))
                        except (TypeError, ValueError):
                            continue
                        text_val = b.get("text")
                        if isinstance(text_val, str):
                            orig_text = b_map.get(bid, "")
                            if is_protected_bullet(orig_text):
                                continue
                            b_map[bid] = text_val

                    new_bullets = []
                    for i, orig_text in enumerate(orig_bullets):
                        idx = i + 1
                        new_text = b_map.get(idx, orig_text)
                        new_bullets.append(new_text)

                    job_new["bullets"] = new_bullets

                new_exp.append(job_new)

            model["experience"] = new_exp


    return model


def scrub_forbidden_terms_from_model(
    model: Dict[str, Any],
    forbidden_terms: List[str],
) -> Dict[str, Any]:
    """
    NO-OP SCRUBBER
    """
    return model


# ============================================================
# ====================== ATS KEYWORDS ========================
# ============================================================

ATS_STOPWORDS = {
    "the",
    "and",
    "or",
    "for",
    "with",
    "a",
    "an",
    "to",
    "of",
    "in",
    "on",
    "at",
    "by",
    "as",
    "is",
    "are",
    "be",
    "will",
    "must",
    "can",
    "should",
    "this",
    "that",
    "you",
    "your",
    "our",
    "we",
}

ATS_SHORT_WHITELIST = {
    "sql",
    "c#",
    "c++",
    "go",
    "r",
    "aws",
    "api",
    "ml",
    "ai",
}

TECH_PHRASE_PATTERNS = [
    r"\bJava\b",
    r"\bSpring Boot\b",
    r"\bSpring\b",
    r"\bREST(?:ful)? APIs?\b",
    r"\bREST APIs?\b",
    r"\bWeb API\b",
    r"\bASP\.NET\b",
    r"\b\.NET\b",
    r"\bC\#\b",
    r"\bC\+\+\b",
    r"\bPython\b",
    r"\bJavaScript\b",
    r"\bTypeScript\b",
    r"\bNode\.js\b",
    r"\bReact\b",
    r"\bAngular\b",
    r"\bVue\b",
    r"\bKubernetes\b",
    r"\bDocker\b",
    r"\bMicroservices?\b",
    r"\bSQL Server(?: \d{4})?\b",
    r"\bOracle\b",
    r"\bPostgreSQL\b",
    r"\bMongoDB\b",
    r"\bMySQL\b",
    r"\bSnowflake\b",
    r"\bDatabricks\b",
    r"\bKafka\b",
    r"\bSpark\b",
    r"\bHadoop\b",
    r"\bSSIS\b",
    r"\bSSRS\b",
    r"\bPower BI\b",
    r"\bAWS\b",
    r"\bAzure\b",
    r"\bGCP\b",
    r"\bS3\b",
    r"\bEC2\b",
    r"\bLambda\b",
    r"\bX12\b",
    r"\bEDI\b",
    r"\bHIPAA\b",
]


def extract_text_from_docx_bytes(data: bytes) -> str:
    """
    Extract plain text (paragraphs + tables) from a DOCX.
    This is ONLY for ATS analysis and does not affect formatting logic.
    """
    file_obj = BytesIO(data)
    doc = Document(file_obj)

    parts: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)

    return "\n".join(parts)




def normalize_text_for_match(text: str) -> str:
    """
    Normalize text for more robust keyword matching:
    - lowercase,
    - keep letters, digits, '+', '#',
    - collapse whitespace,
    - pad with spaces at ends so we can do ' in ' containment checks.
    """
    text = text.lower()
    text = re.sub(r"[^a-z0-9\+\#]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return f" {text} "


def normalize_phrase(phrase: str) -> str:
    """
    Normalize a keyword/phrase to the same form as normalize_text_for_match
    but without padding spaces.
    """
    phrase = phrase.lower()
    phrase = re.sub(r"[^a-z0-9\+\#]+", " ", phrase)
    phrase = re.sub(r"\s+", " ", phrase).strip()
    return phrase


def heuristic_keywords_from_jd(jd: str) -> List[str]:
    """
    Simple heuristic extractor in case GPT fails.
    """
    tokens = re.split(r"[^A-Za-z0-9\+\#]+", jd)
    tokens = [t for t in tokens if t]

    keywords: List[str] = []

    for tok in tokens:
        low = tok.lower()
        if len(low) < 3 and low not in ATS_SHORT_WHITELIST:
            continue
        if low in ATS_STOPWORDS:
            continue
        if any(c.isdigit() for c in tok) or "+" in tok or "#" in tok:
            keywords.append(tok)
        else:
            keywords.append(tok)

    seen = set()
    result = []
    for k in keywords:
        key = k.lower()
        if key not in seen:
            seen.add(key)
            result.append(k)

    return result[:80]


def extract_keywords_with_gpt(
    jd: str,
    llm_provider: str = "openai",
    llm_model: Optional[str] = None,
) -> List[str]:
    """
    Ask LLM to extract technical/domain keywords from the JD AND
    add important implied keywords.
    """
    instructions = (
        "You are helping with ATS keyword analysis.\n\n"
        "The user will provide a job description (JD). It may contain extra sections like:\n"
        "- Company overview\n"
        "- Benefits\n"
        "- Legal/EEO statements\n"
        "- Generic HR text\n\n"
        "Ignore all of that. Focus ONLY on content that describes the role, responsibilities,\n"
        "requirements, and technical environment.\n\n"
        "Your task:\n"
        "1) Identify the IMPORTANT technical and domain-specific keywords that an ATS would look for.\n"
        "2) Also include important implied keywords that are strongly associated with this kind of role.\n"
        "3) Return them as JSON: {\"keywords\": [\"kw1\", \"kw2\", \"kw3\", ...]}.\n"
        "   - Max 80 keywords.\n"
        "   - Deduplicate, case-insensitive.\n"
        "   - Focus on concrete technologies, frameworks, tools, standards, and domain terms.\n"
        "   - Avoid generic verbs or soft skills.\n"
    )

    user_content = f"Job Description:\n\n{jd}"

    resp = get_chat_completion(
        messages=[
            {"role": "system", "content": instructions},
            {"role": "user", "content": user_content},
        ],
        model=llm_model,
        temperature=0.0,
        json_mode=True,
        provider=llm_provider,
    )

    try:
        data = json.loads(resp)
    except json.JSONDecodeError:
        # Fall back to heuristic if the LLM output wasn't valid JSON
        return heuristic_keywords_from_jd(jd)

    kw = data.get("keywords", [])
    keywords: List[str] = []
    for item in kw:
        if isinstance(item, str):
            s = item.strip()
            if s:
                keywords.append(s)
    seen = set()
    result = []
    for k in keywords:
        low = k.lower()
        if low not in seen:
            seen.add(low)
            result.append(k)
    return result[:120]


def suggest_extra_role_keywords(
    jd: str,
    base_keywords: List[str],
    resume_text: str,
    max_count: int = 30,
    llm_provider: str = "openai",
    llm_model: Optional[str] = None,
) -> List[str]:
    """
    Given a JD and the base ATS keywords, ask the LLM for a small set of
    EXTRA role/domain keywords that are strongly related to this role
    (same domain, same seniority), not random tech from other worlds.

    Then filter out anything already in the JD, in the resume, or already
    in the base keyword list.
    """
    jd_clean = (jd or "").strip()
    if not jd_clean:
        return []

    # Short preview of existing keywords so the model doesn't just repeat them
    base_kw_preview = ", ".join(base_keywords[:60])

    system_msg = (
        "You are an ATS and recruiting assistant. "
        "Given a job description and an existing list of ATS keywords, "
        "you suggest a small number of EXTRA role-specific keywords that "
        "are strongly related to this exact role and domain.\n\n"
        "Rules:\n"
        "- Stay in the SAME domain and seniority level as the JD.\n"
        "- Think of things a hiring manager for THIS SPECIFIC JD might also care about,\n"
        "  even if they forgot to write them.\n"
        "- Do NOT propose technologies from unrelated domains or opposite roles.\n"
        "- Return ONLY compact skill/tech/domain keywords (1–3 words each).\n"
        "- Output JSON: {\"extra_keywords\": [\"kw1\", \"kw2\", ...]}."
    )

    user_msg = (
        "JOB DESCRIPTION:\n"
        "----------------\n"
        f"{jd_clean}\n\n"
        "Existing ATS keywords (from the JD and implied):\n"
        f"{base_kw_preview}\n\n"
        "Now propose additional, closely related ATS keywords or tools that are common for roles like this, "
        "but that are NOT simple repeats of the list above. These should be reasonable to add to a resume "
        "for this role (no opposite-domain tech)."
    )

    # Our llm_client.get_chat_completion returns a STRING, not a raw OpenAI object
    raw = get_chat_completion(
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        model=llm_model,
        temperature=0.0,
        json_mode=True,
        provider=llm_provider,
    )

    extra_list: List[str] = []

    # Try to parse JSON first
    try:
        data = json.loads(raw)
        raw_list = data.get("extra_keywords") or data.get("keywords") or []
        if isinstance(raw_list, list):
            for item in raw_list:
                s = str(item).strip()
                if s:
                    extra_list.append(s)
    except Exception:
        # Fallback: split on commas / newlines if JSON parsing fails
        for part in re.split(r"[,\\n]", raw):
            s = part.strip(" \t\r\n-*•")
            if s:
                extra_list.append(s)

    jd_low = jd_clean.lower()
    resume_low = (resume_text or "").lower()
    base_set = {str(k).strip().lower() for k in base_keywords}

    seen: set[str] = set()
    final: List[str] = []
    for kw in extra_list:
        low = kw.lower()
        if not low or low in seen:
            continue
        if low in base_set:
            continue
        if low in jd_low or low in resume_low:
            continue
        # Avoid super long phrases (likely theory, not stack)
        if len(kw.split()) > 5:
            continue
        seen.add(low)
        final.append(kw)
        if len(final) >= max_count:
            break

    return final


def compute_ats_for_text(
    jd: str,
    resume_text: str,
    keywords: Optional[List[str]] = None,
    llm_provider: str = "openai",
    llm_model: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Core ATS computation given JD text and resume plain text.

    If 'keywords' is provided, we use that exact list (for stable before/after
    comparison). If not, we extract keywords from the JD with GPT/heuristic.
    """
    if keywords is None:
        try:
            keywords = extract_keywords_with_gpt(
                jd,
                llm_provider=llm_provider,
                llm_model=llm_model,
            )
        except Exception:
            keywords = heuristic_keywords_from_jd(jd)

    if not keywords:
        return {
            "ats_score": 0,
            "keywords": [],
            "matched_keywords": [],
            "missing_keywords": [],
        }

    # IMPORTANT: use the actual helper name you already have
    normalized_resume = normalize_text_for_match(resume_text)
    normalized_jd = normalize_text_for_match(jd)

    matched: List[str] = []
    missing: List[str] = []

    # Simple containment match for now
    for kw in keywords:
        s = (kw or "").strip()
        if not s:
            continue
        low = s.lower()
        if low in normalized_resume:
            matched.append(s)
        else:
            missing.append(s)

    # Basic score: percentage of keywords matched
    total = len(keywords)
    score = 0
    if total > 0:
        score = int(round(100.0 * len(matched) / total))

    # Suggest a few extra role keywords for user improvement
    extra_recommended = suggest_extra_role_keywords(
        jd,
        keywords,
        resume_text,
        llm_provider=llm_provider,
        llm_model=llm_model,
    )

    return {
        "ats_score": score,
        "keywords": keywords,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "extra_recommended_keywords": extra_recommended,
        "jd_excerpt": normalized_jd[:2000],
        "resume_excerpt": normalized_resume[:2000],
    }




def model_to_plain_text_for_ats(model: Dict[str, Any]) -> str:
    """
    Flatten our structured resume model into a plain text string for ATS-style
    keyword matching.

    We now include:
      - header.lines
      - summary bullets
      - skills rows
      - experience headers/bullets/env/role/project_description
      - education.lines
      - certifications.lines
      - projects.lines
    """
    parts: List[str] = []

    # Header lines (e.g., name in blue, contact info, title)
    header = model.get("header") or {}
    if isinstance(header, dict):
        lines = header.get("lines") or []
        if isinstance(lines, list):
            for line in lines:
                if isinstance(line, str):
                    txt = line.strip()
                    if txt:
                        parts.append(txt)

    # Summary bullets
    summary = model.get("summary") or {}
    if isinstance(summary, dict):
        bullets = summary.get("bullets") or []
        for b in bullets:
            if isinstance(b, str):
                txt = b.strip()
                if txt:
                    parts.append(txt)
    elif isinstance(summary, list):
        # Legacy list-of-dicts
        for b in summary:
            if isinstance(b, dict):
                txt = (b.get("text") or "").strip()
                if txt:
                    parts.append(txt)

    # Skills table
    skills = model.get("skills") or {}
    if isinstance(skills, dict):
        rows = skills.get("rows") or []
        if isinstance(rows, list):
            for row in rows:
                label = None
                value = None

                if isinstance(row, list) and len(row) >= 2:
                    label, value = row[0], row[1]
                elif isinstance(row, dict):
                    label = row.get("label")
                    value = row.get("value")

                if isinstance(label, str):
                    lab = label.strip()
                    if lab:
                        parts.append(lab)
                if isinstance(value, str):
                    val = value.strip()
                    if val:
                        parts.append(val)

    # Experience blocks
    experience = model.get("experience") or []
    if isinstance(experience, list):
        for job in experience:
            if not isinstance(job, dict):
                continue
            header_text = (job.get("header") or "").strip()
            if header_text:
                parts.append(header_text)

            # Label/value style fields
            for fld in ("environment", "role", "project_description"):
                fld_obj = job.get(fld)
                if isinstance(fld_obj, dict):
                    value = (fld_obj.get("value") or "").strip()
                    if value:
                        parts.append(value)

            bullets = job.get("bullets") or []
            if isinstance(bullets, list):
                for b in bullets:
                    if not isinstance(b, str):
                        continue
                    txt = b.strip()
                    if txt:
                        parts.append(txt)

    # Education lines
    education = model.get("education") or {}
    if isinstance(education, dict):
        lines = education.get("lines") or []
        if isinstance(lines, list):
            for line in lines:
                if isinstance(line, str):
                    txt = line.strip()
                    if txt:
                        parts.append(txt)

    # Certifications lines
    certs = model.get("certifications") or {}
    if isinstance(certs, dict):
        lines = certs.get("lines") or []
        if isinstance(lines, list):
            for line in lines:
                if isinstance(line, str):
                    txt = line.strip()
                    if txt:
                        parts.append(txt)

    # Projects lines
    projects = model.get("projects") or {}
    if isinstance(projects, dict):
        lines = projects.get("lines") or []
        if isinstance(lines, list):
            for line in lines:
                if isinstance(line, str):
                    txt = line.strip()
                    if txt:
                        parts.append(txt)

    return "\n".join(parts)


def compute_missing_keywords_for_model(
    jd: str,
    model: Dict[str, Any],
    llm_provider: str = "openai",
    llm_model: Optional[str] = None,
) -> List[str]:
    """
    Compute which ATS keywords from the JD are currently missing from the
    given resume model. This is used when the user chooses 'force missing
    keywords' so we can ask GPT to focus on those gaps.
    """
    resume_text = model_to_plain_text_for_ats(model)

    try:
        keywords = extract_keywords_with_gpt(jd, llm_provider=llm_provider, llm_model=llm_model)
    except Exception:
        keywords = heuristic_keywords_from_jd(jd)

    if not keywords:
        return []

    text_low = resume_text.lower()
    missing: List[str] = []

    for kw in keywords:
        s = (kw or "").strip()
        if not s:
            continue
        if s.lower() not in text_low:
            missing.append(s)

    # Deduplicate while preserving order
    seen = set()
    deduped: List[str] = []
    for s in missing:
        low = s.lower()
        if low not in seen:
            seen.add(low)
            deduped.append(s)

    return deduped



# ============================================================
# ====== POST-PROCESSORS: BOLD COVERAGE & SIMILARITY =========
# ============================================================

STOP_VERBS = {
    "designed",
    "built",
    "implemented",
    "led",
    "optimized",
    "automated",
    "improved",
    "developed",
    "delivered",
    "migrated",
    "refactored",
    "created",
    "configured",
    "tuned",
    "supported",
    "ran",
    "run",
    "managed",
    "maintained",
    "wrote",
    "owned",
    "enhanced",
    "modernized",
    "architected",
    "deployed",
    "integrated",
    "analyzed",
}


def _boldify_first_useful_token(text: str) -> str:
    """
    Aggressively bold known technical phrases using TECH_PHRASE_PATTERNS.

    - If the text already contains '**', we leave it untouched.
    - Otherwise, wrap every non-overlapping match of known technical phrases.
    """
    if not isinstance(text, str):
        return text

    # If GPT or the original resume already added bold, respect it.
    if "**" in text:
        return text

    s = text
    matches = []

    # Collect all matches (start, end) for all tech phrase patterns
    for pattern in TECH_PHRASE_PATTERNS:
        for m in re.finditer(pattern, s):
            start, end = m.span()
            matches.append((start, end))

    if not matches:
        # No known tech phrase found → leave line as-is
        return text

    # Sort and merge overlapping ranges so we don't double-wrap
    matches.sort(key=lambda x: x[0])
    merged = []
    cur_start, cur_end = matches[0]
    for start, end in matches[1:]:
        if start <= cur_end:  # overlap or touch
            cur_end = max(cur_end, end)
        else:
            merged.append((cur_start, cur_end))
            cur_start, cur_end = start, end
    merged.append((cur_start, cur_end))

    # Build new string with ** around each merged range
    out_parts = []
    last_idx = 0
    for start, end in merged:
        # text before the phrase
        if start > last_idx:
            out_parts.append(s[last_idx:start])
        phrase = s[start:end]
        out_parts.append(f"**{phrase}**")
        last_idx = end

    # tail after last match
    if last_idx < len(s):
        out_parts.append(s[last_idx:])

    return "".join(out_parts)


def _strip_markdown_bold(text: Any) -> Any:
    """
    Remove any existing **...** markers but keep the inner text.
    This lets us ignore GPT's own bolding and control it ourselves.
    """
    if not isinstance(text, str):
        return text
    # Replace **something** with just something
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)


def strip_bold_from_model(model: Dict[str, Any]) -> Dict[str, Any]:
    """
    Walk the model and strip **...** from:
    - skills.rows[i].label / value
    - experience[j].header
    - experience[j].environment.value / role.value / project_description.value

    IMPORTANT:
    - We keep any bold markers in summary bullets and experience bullets.
    """
    m = deepcopy(model)

    # --- SUMMARY: KEEP GPT BOLD AS-IS ---

    # --- SKILLS: strip bold from labels/values ---
    skills = m.get("skills")
    if isinstance(skills, dict):
        rows = skills.get("rows")
        if isinstance(rows, list):
            for row in rows:
                # Legacy dict shape
                if isinstance(row, dict):
                    if isinstance(row.get("label"), str):
                        row["label"] = _strip_markdown_bold(row["label"])
                    if isinstance(row.get("value"), str):
                        row["value"] = _strip_markdown_bold(row["value"])
                # Table/list shape: ["Label", "Value"]
                elif isinstance(row, list) and len(row) >= 2:
                    if isinstance(row[0], str):
                        row[0] = _strip_markdown_bold(row[0])
                    if isinstance(row[1], str):
                        row[1] = _strip_markdown_bold(row[1])

    # --- EXPERIENCE: strip bold from header + env/role/project_description ONLY ---
    exp = m.get("experience")
    if isinstance(exp, list):
        for job in exp:
            if not isinstance(job, dict):
                continue

            # Header
            if isinstance(job.get("header"), str):
                job["header"] = _strip_markdown_bold(job["header"])

            # environment / role / project_description
            for fld in ("environment", "role", "project_description"):
                fld_obj = job.get(fld)
                if isinstance(fld_obj, dict) and isinstance(fld_obj.get("value"), str):
                    fld_obj["value"] = _strip_markdown_bold(fld_obj["value"])

            # Bullets: left as-is (we keep GPT bold)

    return m


def ensure_bold_coverage(model: Dict[str, Any]) -> Dict[str, Any]:
    """
    Ensure that:
    - Every summary bullet has at least one **...**,
    - Every non-protected experience bullet has at least one **...**.
    """
    # Summary
    summary = model.get("summary")
    if isinstance(summary, dict):
        bullets = summary.get("bullets")
        if isinstance(bullets, list):
            summary["bullets"] = [
                _boldify_first_useful_token(b) if isinstance(b, str) else b
                for b in bullets
            ]
    elif isinstance(summary, list):
        for item in summary:
            if not isinstance(item, dict):
                continue
            txt = item.get("text")
            if isinstance(txt, str):
                item["text"] = _boldify_first_useful_token(txt)

    # Experience bullets
    exp = model.get("experience")
    if isinstance(exp, list):
        for job in exp:
            if not isinstance(job, dict):
                continue
            bullets = job.get("bullets")
            if not isinstance(bullets, list):
                continue

            new_bullets = []
            for b in bullets:
                if not isinstance(b, str):
                    new_bullets.append(b)
                    continue
                if is_protected_bullet(b):
                    new_bullets.append(b)
                    continue
                new_bullets.append(_boldify_first_useful_token(b))
            job["bullets"] = new_bullets

    return model


def _fix_semicolons_in_bullet(text: Any) -> Any:
    """
    Replace semicolons in bullets with ', and ' in a safe-ish way.
    """
    if not isinstance(text, str):
        return text
    if ";" not in text:
        return text

    s = text

    # If GPT wrote '; and ...', collapse to a single ', and ...'
    import re as _re
    s = _re.sub(r";\s*and\s+", ", and ", s, flags=_re.IGNORECASE)

    # For any remaining ';', turn into ', and '
    s = _re.sub(r";\s*", ", and ", s)

    return s


def cleanup_semicolons_in_model(model: Dict[str, Any]) -> Dict[str, Any]:
    """
    Walk summary + experience bullets and remove semicolons
    using _fix_semicolons_in_bullet.

    - We skip 'protected' heading-style bullets.
    """
    new_model = deepcopy(model)

    # Summary bullets
    summary = new_model.get("summary")
    if isinstance(summary, dict):
        bullets = summary.get("bullets")
        if isinstance(bullets, list):
            summary["bullets"] = [
                _fix_semicolons_in_bullet(b) if isinstance(b, str) else b
                for b in bullets
            ]
        new_model["summary"] = summary

    # Experience bullets
    exp = new_model.get("experience")
    if isinstance(exp, list):
        for job in exp:
            if not isinstance(job, dict):
                continue
            bullets = job.get("bullets")
            if not isinstance(bullets, list):
                continue
            cleaned = []
            for b in bullets:
                if isinstance(b, str) and not is_protected_bullet(b):
                    cleaned.append(_fix_semicolons_in_bullet(b))
                else:
                    cleaned.append(b)
            job["bullets"] = cleaned

    return new_model


def compute_model_similarity(orig_model: Dict[str, Any], new_model: Dict[str, Any]) -> float:
    """
    Compute a rough similarity score between 0 and 1 based on token Jaccard.
    """

    def tokens_from_text(s: str) -> set:
        s = s.lower()
        raw = re.split(r"[^a-z0-9\+\#]+", s)
        toks = set()
        for t in raw:
            if not t:
                continue
            if len(t) < 3 and t not in ATS_SHORT_WHITELIST:
                continue
            if t in ATS_STOPWORDS:
                continue
            toks.add(t)
        return toks

    t_orig = tokens_from_text(model_to_plain_text_for_ats(orig_model))
    t_new = tokens_from_text(model_to_plain_text_for_ats(new_model))

    if not t_orig and not t_new:
        return 0.0
    union = t_orig | t_new
    if not union:
        return 0.0
    inter = t_orig & t_new
    return len(inter) / len(union)





# ============================================================
# =============== RESUME QUALITY / ISSUE DETECTION ===========
# ============================================================

_WEAK_VERB_PREFIXES = [
    "responsible for",
    "worked on",
    "helped with",
    "involved in",
    "participated in",
    "tasked with",
    "duties included",
]

def _word_count(text: str) -> int:
    if not isinstance(text, str):
        return 0
    return len([w for w in text.strip().split() if w])


def _collect_summary_bullets(model: Dict[str, Any]) -> List[str]:
    out: List[str] = []
    summary = model.get("summary") or {}
    if isinstance(summary, dict):
        bullets = summary.get("bullets") or []
        for b in bullets:
            if isinstance(b, str):
                t = b.strip()
                if t:
                    out.append(t)
    elif isinstance(summary, list):
        for item in summary:
            if isinstance(item, dict):
                t = (item.get("text") or "").strip()
                if t:
                    out.append(t)
    return out


def _collect_experience_bullets(model: Dict[str, Any]) -> List[tuple[str, str]]:
    """
    Return list of (item_id, bullet_text) where item_id is like 'exp_1_bullet_3'.
    """
    out: List[tuple[str, str]] = []
    exp = model.get("experience") or []
    if not isinstance(exp, list):
        return out

    for j, job in enumerate(exp, start=1):
        if not isinstance(job, dict):
            continue
        bullets = job.get("bullets") or []
        if not isinstance(bullets, list):
            continue
        for i, b in enumerate(bullets, start=1):
            if not isinstance(b, str):
                continue
            t = b.strip()
            if not t:
                continue
            item_id = f"exp_{j}_bullet_{i}"
            out.append((item_id, t))
    return out


def _detect_weak_verbs(summary_bullets: List[str], exp_bullets: List[tuple[str, str]]) -> List[DetectionIssue]:
    issues: List[DetectionIssue] = []

    # Summary bullets
    for idx, text in enumerate(summary_bullets, start=1):
        tlow = text.strip().lower()
        if any(tlow.startswith(prefix) for prefix in _WEAK_VERB_PREFIXES):
            issues.append(
                DetectionIssue(
                    id=f"summary_{idx}_weak_verb",
                    section="summary",
                    item_id=f"summary_{idx}",
                    severity="warning",
                    code="WEAK_VERB",
                    message="Summary bullet starts with a weak or generic phrase.",
                    suggestion="Rewrite this bullet to start with a concrete action verb (e.g. Designed, Implemented, Built, Automated).",
                )
            )

    # Experience bullets
    for item_id, text in exp_bullets:
        tlow = text.strip().lower()
        if any(tlow.startswith(prefix) for prefix in _WEAK_VERB_PREFIXES):
            issues.append(
                DetectionIssue(
                    id=f"{item_id}_weak_verb",
                    section="experience",
                    item_id=item_id,
                    severity="warning",
                    code="WEAK_VERB",
                    message="Experience bullet starts with a weak or generic phrase.",
                    suggestion="Rewrite this bullet to start with a concrete action verb that reflects what you actually did.",
                )
            )

    return issues


def _detect_bullet_length(summary_bullets: List[str], exp_bullets: List[tuple[str, str]]) -> tuple[List[DetectionIssue], int, int]:
    issues: List[DetectionIssue] = []
    long_count = 0
    short_count = 0

    for idx, text in enumerate(summary_bullets, start=1):
        wc = _word_count(text)
        if wc > 45:
            long_count += 1
            issues.append(
                DetectionIssue(
                    id=f"summary_{idx}_too_long",
                    section="summary",
                    item_id=f"summary_{idx}",
                    severity="info",
                    code="BULLET_TOO_LONG",
                    message=f"Summary bullet is very long ({wc} words).",
                    suggestion="Consider splitting this into two shorter bullets or tightening the wording.",
                )
            )
        elif wc > 0 and wc < 5:
            short_count += 1
            issues.append(
                DetectionIssue(
                    id=f"summary_{idx}_too_short",
                    section="summary",
                    item_id=f"summary_{idx}",
                    severity="info",
                    code="BULLET_TOO_SHORT",
                    message=f"Summary bullet is extremely short ({wc} words).",
                    suggestion="Add a bit more detail so the bullet has a clear action + result.",
                )
            )

    for item_id, text in exp_bullets:
        wc = _word_count(text)
        if wc > 45:
            long_count += 1
            issues.append(
                DetectionIssue(
                    id=f"{item_id}_too_long",
                    section="experience",
                    item_id=item_id,
                    severity="info",
                    code="BULLET_TOO_LONG",
                    message=f"Experience bullet is very long ({wc} words).",
                    suggestion="Consider splitting into two bullets or tightening the sentence.",
                )
            )
        elif wc > 0 and wc < 5:
            short_count += 1
            issues.append(
                DetectionIssue(
                    id=f"{item_id}_too_short",
                    section="experience",
                    item_id=item_id,
                    severity="info",
                    code="BULLET_TOO_SHORT",
                    message=f"Experience bullet is extremely short ({wc} words).",
                    suggestion="Add more context so the bullet shows what you did and why it mattered.",
                )
            )

    return issues, long_count, short_count


def _detect_summary_duplication(model: Dict[str, Any], summary_bullets: List[str]) -> List[DetectionIssue]:
    issues: List[DetectionIssue] = []
    if not summary_bullets:
        return issues

    certs = model.get("certifications") or {}
    cert_text = ""
    if isinstance(certs, dict):
        lines = certs.get("lines") or []
        if isinstance(lines, list):
            cert_text = "\n".join([str(x) for x in lines if isinstance(x, str)]).lower()

    if not cert_text.strip():
        return issues

    for idx, bullet in enumerate(summary_bullets, start=1):
        b = bullet.strip()
        if len(b) < 20:
            continue
        snippet = b.lower()[:80]
        if snippet and snippet in cert_text:
            issues.append(
                DetectionIssue(
                    id=f"summary_{idx}_duplicated_in_certs",
                    section="global",
                    item_id=f"summary_{idx}",
                    severity="warning",
                    code="SUMMARY_DUPLICATED",
                    message="A summary bullet appears to be duplicated inside Certifications or another section.",
                    suggestion="Avoid copying summary bullets into other sections; keep Certifications focused on actual credentials.",
                )
            )
    return issues


def _detect_missing_sections(model: Dict[str, Any]) -> List[DetectionIssue]:
    issues: List[DetectionIssue] = []

    # Summary
    summary_ok = False
    summary = model.get("summary")
    if isinstance(summary, dict):
        bullets = summary.get("bullets") or []
        if any(isinstance(b, str) and b.strip() for b in bullets):
            summary_ok = True

    if not summary_ok:
        issues.append(
            DetectionIssue(
                id="missing_summary",
                section="summary",
                item_id=None,
                severity="warning",
                code="MISSING_SUMMARY",
                message="No clear professional summary was detected.",
                suggestion="Add 2–4 summary bullets that highlight your domain, years of experience, and strongest skills.",
            )
        )

    # Skills
    skills_ok = False
    skills = model.get("skills")
    if isinstance(skills, dict):
        rows = skills.get("rows") or []
        if any(row for row in rows):
            skills_ok = True

    if not skills_ok:
        issues.append(
            DetectionIssue(
                id="missing_skills",
                section="skills",
                item_id=None,
                severity="error",
                code="MISSING_SKILLS",
                message="No skills table was detected.",
                suggestion="Add a skills section with grouped tools/technologies that match the target job.",
            )
        )

    # Experience
    exp_ok = False
    exp = model.get("experience")
    if isinstance(exp, list) and exp:
        exp_ok = True

    if not exp_ok:
        issues.append(
            DetectionIssue(
                id="missing_experience",
                section="experience",
                item_id=None,
                severity="error",
                code="MISSING_EXPERIENCE",
                message="No work experience section was detected.",
                suggestion="Add at least one recent role with bullets describing what you built or owned.",
            )
        )

    return issues


def _compute_keyword_alignment_score(jd: str, resume_text: str) -> tuple[int, List[DetectionIssue]]:
    """
    Cheap, non-LLM keyword alignment:
    - Use heuristic_keywords_from_jd to get keywords.
    - Score = % of those that appear in resume_text.
    """
    issues: List[DetectionIssue] = []
    jd_clean = (jd or "").strip()
    if not jd_clean:
        return 0, issues

    try:
        keywords = heuristic_keywords_from_jd(jd_clean)
    except Exception:
        keywords = []

    if not keywords:
        return 0, issues

    res_low = resume_text.lower()
    matched = 0
    for kw in keywords:
        s = (kw or "").strip()
        if not s:
            continue
        if s.lower() in res_low:
            matched += 1

    total = len(keywords)
    score = int(round(100.0 * matched / total)) if total > 0 else 0

    if score < 40:
        issues.append(
            DetectionIssue(
                id="keyword_alignment_low",
                section="global",
                item_id=None,
                severity="warning",
                code="KEYWORD_ALIGNMENT_LOW",
                message=f"Only about {score}% of important JD keywords appear in the resume.",
                suggestion="Make sure your summary, skills, and experience explicitly mention the main tools and concepts from the job description.",
            )
        )

    return score, issues


def run_detection_engine(model: Dict[str, Any], job_description: Optional[str] = None) -> DetectionReport:
    """
    Pure-Python detection of common issues:
    - weak verbs
    - bullet length extremes
    - missing sections
    - summary duplicated in certifications
    - JD keyword alignment (cheap heuristic)
    """
    issues: List[DetectionIssue] = []

    # For potential future use: template-level metadata
    template_case = None
    tmpl = model.get("template_meta") or {}
    if isinstance(tmpl, dict):
        tc = tmpl.get("resume_template")
        if isinstance(tc, str) and tc.strip():
            template_case = tc.strip()

    summary_bullets = _collect_summary_bullets(model)
    exp_bullets = _collect_experience_bullets(model)

    # Weak verbs
    weak_issues = _detect_weak_verbs(summary_bullets, exp_bullets)
    issues.extend(weak_issues)
    weak_count = len(weak_issues)

    # Bullet length
    length_issues, long_count, short_count = _detect_bullet_length(summary_bullets, exp_bullets)
    issues.extend(length_issues)

    # Missing sections
    issues.extend(_detect_missing_sections(model))

    # Summary duplication into Certifications
    issues.extend(_detect_summary_duplication(model, summary_bullets))

    # Keyword alignment if JD provided
    resume_text = model_to_plain_text_for_ats(model)
    keyword_alignment = 0
    if job_description:
        ka_score, ka_issues = _compute_keyword_alignment_score(job_description, resume_text)
        keyword_alignment = ka_score
        issues.extend(ka_issues)

    # ---- Scores ----
    # Bullet strength: penalize weak verbs + extremes
    bullet_strength = 100
    bullet_strength -= weak_count * 5
    bullet_strength -= long_count * 3
    bullet_strength -= short_count * 2
    bullet_strength = max(0, min(100, bullet_strength))

    # Clarity: mostly length-based
    clarity = 100
    clarity -= long_count * 4
    clarity -= short_count * 1
    clarity = max(0, min(100, clarity))

    # Overall quality as a blend
    overall_quality = int(
        round(0.4 * bullet_strength + 0.3 * clarity + 0.3 * keyword_alignment)
    )

    scores = DetectionScores(
        overall_quality=overall_quality,
        bullet_strength=bullet_strength,
        keyword_alignment=keyword_alignment,
        clarity=clarity,
    )

    return DetectionReport(
        template_case=template_case,
        scores=scores,
        issues=issues,
    )


@app.post("/api/detect-model")
def detect_model(req: DetectRequest):
    """
    Run heuristic quality/template checks on the current resume model.

    Input:
      - model_json: the same JSON model used on Page 2.
      - job_description: optional JD text (improves keyword alignment score).
    """
    try:
        model = json.loads(req.model_json)
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=400,
            detail="Invalid JSON in model_json.",
        )

    if not isinstance(model, dict):
        raise HTTPException(
            status_code=400,
            detail="model_json must encode a JSON object.",
        )

    report = run_detection_engine(model, req.job_description)
    # FastAPI will auto-serialize the Pydantic model
    return report




# ============================================================
# ==================== AUTH & PASSWORD RESET =================
# ============================================================

@app.post("/auth/register")
def register(req: RegisterRequest, db: Session = Depends(get_db)):
    """
    Register with email + password.

    Cases:
    - If a normal account already exists (email + password_hash): block with error.
    - If a Google-only account exists (email + google_id, no password_hash): attach a password.
    - If no account exists: create a new local account.
    """
    existing = db.query(User).filter(User.email == req.email).first()

    if existing:
        if existing.password_hash:
            # Real local account already exists for this email
            raise HTTPException(status_code=400, detail="Email already registered")

        # Google-only account exists: upgrade by setting a password
        existing.password_hash = pbkdf2_sha256.hash(req.password)
        db.commit()
        db.refresh(existing)

        token = create_token(existing.id)
        return AuthResponse(token=token, email=existing.email, user_id=existing.id)

    # No user exists yet: create a brand new local account
    password_hash = pbkdf2_sha256.hash(req.password)
    user = User(email=req.email, password_hash=password_hash)
    db.add(user)
    db.commit()
    db.refresh(user)

    token = create_token(user.id)
    return AuthResponse(token=token, email=user.email, user_id=user.id)



@app.post("/auth/login", response_model=AuthResponse)
def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == req.email).first()
    if not user or not user.password_hash:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid credentials",
        )

    if not pbkdf2_sha256.verify(req.password, user.password_hash):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid credentials",
        )

    token = create_token(user.id)
    return AuthResponse(token=token, email=user.email, user_id=user.id)




@app.get("/auth/google/start")
def google_start():
    params = {
        "client_id": GOOGLE_CLIENT_ID,
        "redirect_uri": GOOGLE_REDIRECT_URI,
        "response_type": "code",
        "scope": "openid email profile",
        "access_type": "offline",
        "prompt": "consent",
    }
    url = "https://accounts.google.com/o/oauth2/v2/auth?" + urlencode(params)
    return RedirectResponse(url)


@app.get("/auth/google/callback")
def google_callback(code: str, db: Session = Depends(get_db)):
    # Exchange code for tokens
    token_resp = httpx.post(
        "https://oauth2.googleapis.com/token",
        data={
            "code": code,
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "redirect_uri": GOOGLE_REDIRECT_URI,
            "grant_type": "authorization_code",
        },
    )
    token_resp.raise_for_status()
    token_data = token_resp.json()

    id_token = token_data["id_token"]

    # Decode ID token (we trust Google, so skip signature verification here in dev)
    google_info = jwt.decode(id_token, options={"verify_signature": False})

    google_id = google_info["sub"]
    email = google_info["email"]

    # Find or create user
    user = db.query(User).filter(User.email == email).first()
    if user:
        if not user.google_id:
            user.google_id = google_id
            db.commit()
            db.refresh(user)
    else:
        user = User(email=email, google_id=google_id)
        db.add(user)
        db.commit()
        db.refresh(user)

    app_token = create_token(user.id)

    # Redirect back to frontend with token in URL fragment
    frontend_url = f"{FRONTEND_BASE_URL}/index.html#token={app_token}"
    return RedirectResponse(frontend_url)


@app.get("/auth/me")
def get_me(current_user: User = Depends(get_current_user)):
    return {
        "user_id": current_user.id,
        "email": current_user.email,
        "has_password": bool(current_user.password_hash),
        "has_google": bool(current_user.google_id),
    }

@app.post("/auth/forgot")
def forgot_password(req: ForgotPasswordRequest, db: Session = Depends(get_db)):
    """
    Step 1: Forgot password endpoint.
    - Takes an email.
    - If it belongs to a local (password) user, create a reset token.
    - Send the reset link by email via SMTP.
    - Always return {"status": "ok"} so we don't leak which emails exist.
    """
    # Look up user by email
    user = db.query(User).filter(User.email == req.email).first()

    # If user doesn't exist or is Google-only (no password_hash), we still
    # return OK and do nothing, to avoid information leaks.
    if not user or not user.password_hash:
        return {"status": "ok"}

    # Generate a random token string (this is what goes in the URL)
    raw_token = secrets.token_urlsafe(32)

    # Hash it before storing in DB (we never store the raw token)
    token_hash = pbkdf2_sha256.hash(raw_token)

    # Token expires in 1 hour
    expires_at = dt.datetime.utcnow() + dt.timedelta(hours=1)

    reset = PasswordResetToken(
        user_id=user.id,
        token_hash=token_hash,
        expires_at=expires_at,
    )
    db.add(reset)
    db.commit()

    # Build reset URL for the frontend
    # You can later change FRONTEND_BASE_URL to your real deployed URL
    FRONTEND_BASE_URL = os.getenv("FRONTEND_BASE_URL", "http://localhost:5500")
    reset_url = f"{FRONTEND_BASE_URL}/reset-password.html?token={reset_token}"


    # Dev log (optional)
    print(f"[RESET] Password reset link for {user.email}: {reset_url}")

    # Send email via SMTP
    try:
        send_password_reset_email(user.email, reset_url)
    except Exception as e:
        # For now just log the error so it doesn't crash the endpoint
        print(f"[RESET] Error sending reset email to {user.email}: {e}")

    # Frontend will just show a generic 'If this email is registered, we sent a link'
    return {"status": "ok"}



@app.post("/auth/reset-password")
def reset_password(req: ResetPasswordRequest, db: Session = Depends(get_db)):
    """
    Step 2 of password reset:
    - Takes raw token and new password.
    - Finds a matching, unused, unexpired PasswordResetToken.
    - Updates the user's password_hash.
    - Marks the token as used.
    """
    # Basic validation
    if req.new_password != req.confirm_password:
        raise HTTPException(status_code=400, detail="Passwords do not match")

    # Grab all unused reset tokens that haven't yet expired
    now = dt.datetime.utcnow()
    reset_rows = (
        db.query(PasswordResetToken)
        .filter(
            PasswordResetToken.used_at.is_(None),
            PasswordResetToken.expires_at > now,
        )
        .all()
    )

    matched = None
    # We stored a *hash* of the token, so we must verify each
    for row in reset_rows:
        if pbkdf2_sha256.verify(req.token, row.token_hash):
            matched = row
            break

    if not matched:
        # Token invalid or expired
        raise HTTPException(status_code=400, detail="Invalid or expired token")

    # Fetch the user this token belongs to
    user = db.query(User).filter(User.id == matched.user_id).first()
    if not user:
        raise HTTPException(status_code=400, detail="User for this token no longer exists")

    # Update the user's password
    new_hash = pbkdf2_sha256.hash(req.new_password)
    user.password_hash = new_hash

    # Mark this token as used
    matched.used_at = now

    db.commit()

    return {"status": "ok"}



@app.post("/api/ats-from-docx")
async def ats_from_docx(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
):
    """
    ATS analysis on the ORIGINAL resume.
    """
    jd = job_description.strip()
    if not jd:
        return JSONResponse(
            status_code=400, content={"error": "Job description is empty."}
        )

    data = await resume_file.read()
    try:
        resume_text = extract_text_from_docx_bytes(data)
    except Exception as e:
        return JSONResponse(
            status_code=500, content={"error": f"Failed to read DOCX: {e}"}
        )

    result = compute_ats_for_text(jd, resume_text)
    return {"ats": result}


@app.post("/api/ats-from-updated-model")
async def ats_from_updated_model(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    model_json: str = Form(...),
):
    """
    ATS analysis on the UPDATED resume model.
    We reconstruct a DOCX from the model and run ATS on that text.
    """
    jd = job_description.strip()
    if not jd:
        return JSONResponse(
            status_code=400, content={"error": "Job description is empty."}
        )

    data = await resume_file.read()
    try:
        doc = load_document_from_bytes(data)
    except Exception as e:
        return JSONResponse(
            status_code=500, content={"error": f"Failed to load base DOCX: {e}"}
        )

    try:
        model_dict = json.loads(model_json)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400, content={"error": "Invalid JSON for updated model."}
        )

    try:
        updated_bytes = apply_model_and_generate_docx(data, model_dict)
        resume_text = extract_text_from_docx_bytes(updated_bytes)
        result = compute_ats_for_text(jd, resume_text)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to run ATS analysis on updated resume: {e}"},
        )

    return {"ats": result}


# ============================================================
# =========== WRAPPER ATS ENDPOINTS FOR FRONTEND =============
# ============================================================

@app.post("/api/ats-analyze")
async def ats_analyze(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    llm_providers: Optional[str] = Form(None),
):
    """
    Wrapper used by the frontend to compute ATS score for the original resume.
    """
    jd = job_description.strip()
    if not jd:
        return JSONResponse(
            status_code=400,
            content={"error": "Job description is empty."},
        )

    data = await resume_file.read()
    try:
        resume_text = extract_text_from_docx_bytes(data)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to read DOCX: {e}"},
        )

    # Decide provider/model based on the llm_providers form field
    provider_candidates = parse_llm_providers_field(llm_providers)
    provider_choice = choose_llm_from_list(provider_candidates)
    llm_provider = provider_choice["provider"]
    llm_model = provider_choice["model"]

    result = compute_ats_for_text(
        jd,
        resume_text,
        llm_provider=llm_provider,
        llm_model=llm_model,
    )

    return {
        "ats": result,
        "llm_provider": llm_provider,
        "llm_model": llm_model,
    }


@app.post("/api/ats-analyze-updated")
async def ats_analyze_updated(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    model_json: str = Form(...),
    keywords_json: str = Form(None),
    llm_providers: Optional[str] = Form(None),
):
    """
    Wrapper used by the frontend to compute ATS score for the updated resume model.
    """
    jd = job_description.strip()
    if not jd:
        return JSONResponse(
            status_code=400,
            content={"error": "Job description is empty."},
        )

    data = await resume_file.read()
    try:
        # Basic validation that the uploaded file is a readable DOCX
        _doc = load_document_from_bytes(data)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to load base DOCX: {e}"},
        )

    try:
        model_dict = json.loads(model_json)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={"error": "Invalid JSON for updated model."},
        )

    stable_keywords: Optional[List[str]] = None
    if keywords_json:
        try:
            parsed = json.loads(keywords_json)
            if isinstance(parsed, list):
                clean_kw: List[str] = []
                for x in parsed:
                    s = str(x).strip()
                    if s:
                        clean_kw.append(s)
                if clean_kw:
                    stable_keywords = clean_kw
        except json.JSONDecodeError:
            stable_keywords = None

    # Decide provider/model based on the llm_providers form field
    provider_candidates = parse_llm_providers_field(llm_providers)
    provider_choice = choose_llm_from_list(provider_candidates)
    llm_provider = provider_choice["provider"]
    llm_model = provider_choice["model"]

    try:
        # Apply the updated model to the original DOCX bytes
        updated_bytes = apply_model_and_generate_docx(data, model_dict)
        # Extract plain text from the updated DOCX for ATS
        resume_text = extract_text_from_docx_bytes(updated_bytes)
        # Use a stable keyword list if provided
        result = compute_ats_for_text(
            jd,
            resume_text,
            keywords=stable_keywords,
            llm_provider=llm_provider,
            llm_model=llm_model,
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to run ATS analysis on updated resume: {e}"},
        )

    return {
        "ats": result,
        "llm_provider": llm_provider,
        "llm_model": llm_model,
    }


# ============================================================
# =========== UNIFIED SECTION + ALL-SECTIONS REWRITE =========
# ============================================================

@app.post("/api/rewrite-sections")
async def rewrite_sections(
    job_description: str = Form(...),
    section: str = Form(...),
    model_json: str = Form(...),
    include_ats_keywords: str = Form("false"),
    force_missing_keywords: str = Form("false"),
    llm_providers: Optional[str] = Form(None),
):
    """
    Unified rewrite endpoint used by the frontend.

    - section = "summary" | "skills" | "experience"  -> only that section is touched
    - section = "all"                                -> all editable sections rewritten
    """
    sec = (section or "").strip().lower()
    if sec not in ("summary", "skills", "experience", "all"):
        return JSONResponse(
            status_code=400,
            content={"error": "Invalid section. Must be one of: summary, skills, experience, all."},
        )

    jd = (job_description or "").strip()
    if not jd:
        return JSONResponse(
            status_code=400,
            content={"error": "Job description is empty."},
        )

    try:
        original_model: Dict[str, Any] = json.loads(model_json)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={"error": "Invalid JSON in 'model_json' field."},
        )

    # Decide which LLM provider/model to use for this rewrite request
    provider_candidates = parse_llm_providers_field(llm_providers)
    provider_choice = choose_llm_from_list(provider_candidates)
    llm_provider = provider_choice["provider"]
    llm_model = provider_choice["model"]

    # Build base payload from the full model
    payload = build_rewrite_payload(jd, original_model)

    include_ats_flag = str(include_ats_keywords or "").strip().lower() == "true"
    force_missing_flag = str(force_missing_keywords or "").strip().lower() == "true"

    # If the caller explicitly disables ATS, mark it
    if not include_ats_flag:
        payload["disable_ats"] = True

    # If the caller wants to force missing keywords, compute them from the current model
    if force_missing_flag:
        try:
            missing_keywords = compute_missing_keywords_for_model(
                jd,
                original_model,
                llm_provider=llm_provider,
                llm_model=llm_model,
            )
        except Exception:
            missing_keywords = []

        if missing_keywords:
            payload["ats_keywords"] = missing_keywords
            payload["ats_force_all"] = True
        else:
            # Still force ATS coverage, but nothing extra to add
            payload["ats_force_all"] = True

    # Collect forbidden terms (old domain tech not in JD)
    forbidden_terms = collect_forbidden_terms_from_model(jd, original_model)
    if forbidden_terms:
        payload["forbidden_terms"] = forbidden_terms

    # --- SPECIAL CASE: DATA ENGINEER JD ---
    jd_low = jd.lower()
    is_data_engineer_jd = ("data engineer" in jd_low) or ("data engineering" in jd_low)
    if is_data_engineer_jd:
        # Hint to GPT so it forces Data Engineer titles and stack
        payload["target_role_hint"] = "Data Engineer"

        # Force GPT to rewrite environment and role fully for each job
        exp = payload.get("experience")
        if isinstance(exp, list):
            for job in exp:
                env = job.get("environment")
                if isinstance(env, dict):
                    env["value"] = ""  # GPT must fill with Data Engineering tools
                role_obj = job.get("role")
                if isinstance(role_obj, dict):
                    role_obj["value"] = ""  # GPT must write Data Engineer style title

    # Clone payload for a possible “stronger” second run
    payload_stronger = deepcopy(payload)
    payload_stronger["force_stronger"] = True

    try:
        # ---- Attempt 1 ----
        rewrite_json = call_gpt_rewrite(
            payload,
            llm_provider=llm_provider,
            llm_model=llm_model,
        )
        if not isinstance(rewrite_json, dict):
            raise ValueError(
                f"LLM rewrite response is not a JSON object. Got: {type(rewrite_json)}"
            )

        # Only keep the requested section if sec != "all"
        if sec == "all":
            merged_input = rewrite_json
        else:
            filtered_rewrite: Dict[str, Any] = {}
            if "schema_version" in rewrite_json:
                filtered_rewrite["schema_version"] = rewrite_json["schema_version"]
            if sec in rewrite_json:
                filtered_rewrite[sec] = rewrite_json[sec]
            merged_input = filtered_rewrite

        updated_model = merge_rewrite_into_model(original_model, merged_input)
        forbidden_terms_for_scrub = payload.get("forbidden_terms") or []
        updated_model = scrub_forbidden_terms_from_model(
            updated_model,
            forbidden_terms_for_scrub,
        )

        # Similarity check: if too close to original, do a second, stronger run
        sim = compute_model_similarity(original_model, updated_model)
        if sim > 0.8:
            rewrite_json_2 = call_gpt_rewrite(
                payload_stronger,
                llm_provider=llm_provider,
                llm_model=llm_model,
            )
            if isinstance(rewrite_json_2, dict):
                if sec == "all":
                    merged_input_2 = rewrite_json_2
                else:
                    filtered_rewrite_2: Dict[str, Any] = {}
                    if "schema_version" in rewrite_json_2:
                        filtered_rewrite_2["schema_version"] = rewrite_json_2["schema_version"]
                    if sec in rewrite_json_2:
                        filtered_rewrite_2[sec] = rewrite_json_2[sec]
                    merged_input_2 = filtered_rewrite_2

                updated_model_2 = merge_rewrite_into_model(original_model, merged_input_2)
                updated_model_2 = scrub_forbidden_terms_from_model(
                    updated_model_2,
                    forbidden_terms_for_scrub,
                )
                updated_model = updated_model_2

        # Post-processing: bold, semicolons
        updated_model = strip_bold_from_model(updated_model)
        updated_model = ensure_bold_coverage(updated_model)
        updated_model = cleanup_semicolons_in_model(updated_model)

        # ---------- HEADER ROLE SYNC: update header.lines across all header formats ----------
        hdr = updated_model.get("header")
        if isinstance(hdr, dict):
            # figure out desired role from first experience job
            desired_role = None
            exp_list = updated_model.get("experience")
            if isinstance(exp_list, list) and exp_list:
                first_job = exp_list[0]
                if isinstance(first_job, dict):
                    # Prefer explicit Role: value
                    role_obj = first_job.get("role")
                    if isinstance(role_obj, dict):
                        val = (role_obj.get("value") or "").strip()
                        if val:
                            desired_role = val
                    # Fallback: prefix of job header before first comma
                    if not desired_role:
                        htxt = (first_job.get("header") or "").strip()
                        if htxt:
                            desired_role = htxt.split(",", 1)[0].strip()

            lines = hdr.get("lines")
            if desired_role and isinstance(lines, list) and lines:
                # Helper: detect obvious contact / non-title lines
                def _is_contact_like_segment(seg: str) -> bool:
                    if not seg:
                        return False
                    lower = seg.lower()
                    markers = [
                        "@",
                        "linkedin",
                        "http",
                        "www.",
                        "phone",
                        "email",
                        "+1",
                        "contact",
                        ".com",
                    ]
                    if any(m in lower for m in markers):
                        return True
                    digit_count = sum(ch.isdigit() for ch in seg)
                    return digit_count >= 5

                role_keywords = [
                    "developer",
                    "engineer",
                    "analyst",
                    "architect",
                    "consultant",
                    "specialist",
                    "administrator",
                    "lead",
                    "manager",
                ]

                # Flatten all header lines into visual segments with indices
                flat_segments = []  # list of (line_idx, seg_idx, text)
                for li, line in enumerate(lines):
                    if not isinstance(line, str):
                        continue
                    chunks = line.split("\n")
                    for si, raw_seg in enumerate(chunks):
                        seg = (raw_seg or "").strip()
                        if seg:
                            flat_segments.append((li, si, seg))

                if flat_segments:
                    # 1) Try to find a segment that looks like a role/title
                    candidate_idx = None
                    for idx, (li, si, seg) in enumerate(flat_segments):
                        if _is_contact_like_segment(seg):
                            continue
                        low = seg.lower()
                        if any(k in low for k in role_keywords):
                            candidate_idx = idx
                            break

                    # 2) Fallback: use the second visual line if no role-looking segment
                    if candidate_idx is None and len(flat_segments) >= 2:
                        candidate_idx = 1

                    if candidate_idx is not None:
                        target_line_idx, target_seg_idx, _ = flat_segments[candidate_idx]

                        # Rebuild line chunks preserving original newline structure
                        line_chunks = []
                        for li, line in enumerate(lines):
                            if isinstance(line, str):
                                line_chunks.append(line.split("\n"))
                            else:
                                line_chunks.append([""])

                        # Ensure target indices are in range
                        if (
                            0 <= target_line_idx < len(line_chunks)
                            and 0 <= target_seg_idx < len(line_chunks[target_line_idx])
                        ):
                            line_chunks[target_line_idx][target_seg_idx] = desired_role
                            new_lines = ["\n".join(chunks) for chunks in line_chunks]
                            hdr["lines"] = new_lines

            # We don't need a separate 'title' field anymore
            if "title" in hdr:
                hdr.pop("title", None)

    except ValueError as ve:
        return JSONResponse(
            status_code=500,
            content={"error": str(ve)},
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Rewrite failed: {e}"},
        )

    # NEW: compute ATS on the updated model so the frontend can refresh score + keywords automatically
    ats_result = None
    try:
        updated_resume_text = model_to_plain_text_for_ats(updated_model)
        ats_result = compute_ats_for_text(
            jd,
            updated_resume_text,
            llm_provider=llm_provider,
            llm_model=llm_model,
        )
    except Exception:
        ats_result = None

    response_payload: Dict[str, Any] = {"model": updated_model}
    response_payload["llm_provider"] = llm_provider
    response_payload["llm_model"] = llm_model
    if ats_result is not None:
        response_payload["ats"] = ats_result

    return response_payload



@app.post("/api/generate-cover-letter")
async def generate_cover_letter_api(
    job_description: str = Form(...),
    model_json: str = Form(...),
    llm_providers: Optional[str] = Form(None),
):
    """
    Generate a tailored cover letter using the updated resume model + JD.

    - Uses the same multi-LLM selection logic as ATS/rewrite.
    - Does not require uploading any cover-letter file.
    """
    jd = (job_description or "").strip()
    if not jd:
        return JSONResponse(
            status_code=400,
            content={"error": "Job description is empty."},
        )

    try:
        model_dict = json.loads(model_json)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={"error": "Invalid JSON for model."},
        )

    # 👇 EXACTLY the same pattern as ATS/rewrite
    provider_candidates = parse_llm_providers_field(llm_providers)
    provider_choice = choose_llm_from_list(provider_candidates)
    llm_provider = provider_choice["provider"]
    llm_model = provider_choice["model"]

    try:
        cover_letter_text = generate_cover_letter_text(
            model_dict,
            jd,
            llm_provider=llm_provider,
            llm_model=llm_model,
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to generate cover letter: {e}"},
        )

    return {
        "cover_letter": cover_letter_text,
        "llm_provider": llm_provider,
        "llm_model": llm_model,
    }



@app.post("/api/generate-cover-letter-docx")
async def generate_cover_letter_docx_api(
    cover_letter_text: str = Form(...),
    model_json: Optional[str] = Form(None),
):
    """
    Build a DOCX cover letter document from already-generated cover letter text.

    We optionally accept the updated resume model so we can reuse header lines
    (name + contact info) at the top of the letter and build a nice filename.
    """
    text = (cover_letter_text or "").strip()
    if not text:
        return JSONResponse(
            status_code=400,
            content={"error": "Cover letter text is empty."},
        )

    model_dict: Dict[str, Any] = {}
    if model_json:
        try:
            model_dict = json.loads(model_json)
            if not isinstance(model_dict, dict):
                model_dict = {}
        except json.JSONDecodeError:
            # Not fatal – we can still build a DOCX without header info
            model_dict = {}

    try:
        docx_bytes = build_cover_letter_docx(model_dict, text)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to generate cover letter DOCX: {e}"},
        )

    # 🔑 dynamic filename: name + role + cover_letter
    filename = build_filename_from_model(model_dict, kind="cover_letter")

    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )



# ============================================================
# ================= DOCX PARSE / SAVE ENDPOINTS ==============
# ============================================================

@app.post("/api/parse-docx")
async def parse_docx(file: UploadFile = File(...)):
    """
    Parse a DOCX resume into our structured JSON model.
    """
    data = await file.read()
    try:
        doc = load_document_from_bytes(data)
        model = parse_docx_to_model(doc)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to parse DOCX: {e}"},
        )
    return {"model": model}


@app.post("/api/apply-model")
async def apply_model(
    file: UploadFile = File(...),
    model_json: str = Form(...),
):
    """
    Apply a JSON model to the base DOCX and return the updated DOCX.
    """
    data = await file.read()
    try:
        base_doc = load_document_from_bytes(data)
        model_dict = json.loads(model_json)
        updated_bytes = apply_model_and_generate_docx(data, model_dict)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={"error": "Invalid JSON for model."},
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to apply model: {e}"},
        )

    # 🔑 build filename from the model header
    filename = build_filename_from_model(model_dict, kind="resume")

    return StreamingResponse(
        BytesIO(updated_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Take DOCX bytes, call LibreOffice (soffice) in headless mode,
    and return the resulting PDF as bytes.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "input.docx")
        output_dir = tmpdir
        output_path = os.path.join(output_dir, "input.pdf")

        # Write the DOCX to disk
        with open(input_path, "wb") as f:
            f.write(docx_bytes)

        # Call LibreOffice to convert DOCX -> PDF
        cmd = [
            "soffice",          # relies on soffice being in PATH
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path,
        ]

        try:
            result = subprocess.run(
                cmd,
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
        except subprocess.CalledProcessError as e:
            stderr_text = e.stderr.decode(errors="ignore") if e.stderr else ""
            raise RuntimeError(f"LibreOffice conversion failed: {stderr_text}")

        if not os.path.exists(output_path):
            raise RuntimeError("LibreOffice did not produce a PDF file.")

        with open(output_path, "rb") as f:
            return f.read()


def _rename_docx_to_pdf(filename: str) -> str:
    """
    Utility to turn "Something.docx" into "Something.pdf" while preserving
    the base name if the extension is missing or different.
    """
    if not isinstance(filename, str) or not filename:
        return "output.pdf"

    lower = filename.lower()
    if lower.endswith(".docx"):
        return filename[:-5] + ".pdf"
    if lower.endswith(".doc"):
        return filename[:-4] + ".pdf"
    return filename + ".pdf"


@app.post("/download_resume_pdf")
async def download_resume_pdf(
    file: UploadFile = File(...),
    model_json: str = Form(...),
):
    """
    Apply the current resume model to the uploaded DOCX and return a PDF.
    """
    docx_bytes = await file.read()
    try:
        model_dict = json.loads(model_json)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={"error": "Invalid JSON for model."},
        )

    try:
        updated_docx = apply_model_and_generate_docx(docx_bytes, model_dict)
        pdf_bytes = convert_docx_to_pdf_bytes(updated_docx)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to generate PDF: {e}"},
        )

    filename = build_filename_from_model(model_dict, kind="resume")
    pdf_name = _rename_docx_to_pdf(filename)

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{pdf_name}"'},
    )


@app.post("/download_cover_letter_pdf")
async def download_cover_letter_pdf(
    cover_letter_text: str = Form(...),
    model_json: Optional[str] = Form(None),
):
    """
    Build a PDF cover letter using the generated text and optional model header.
    """
    text = (cover_letter_text or "").strip()
    if not text:
        return JSONResponse(
            status_code=400,
            content={"error": "Cover letter text is empty."},
        )

    model_dict: Dict[str, Any] = {}
    if model_json:
        try:
            model_dict = json.loads(model_json)
            if not isinstance(model_dict, dict):
                model_dict = {}
        except json.JSONDecodeError:
            model_dict = {}

    try:
        docx_bytes = build_cover_letter_docx(model_dict, text)
        pdf_bytes = convert_docx_to_pdf_bytes(docx_bytes)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to generate cover letter PDF: {e}"},
        )

    filename = build_filename_from_model(model_dict, kind="cover_letter")
    pdf_name = _rename_docx_to_pdf(filename)

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{pdf_name}"'},
    )


@app.post("/api/preview-docx")
async def preview_docx(
    file: UploadFile = File(...),
    model_json: Optional[str] = Form(None),
):
    """
    Return a PDF preview of either:
      - the original DOCX (if model_json is None), or
      - the DOCX after applying the current model (if model_json is provided).
    """
    docx_bytes = await file.read()

    if model_json:
        try:
            model = json.loads(model_json)
        except json.JSONDecodeError as e:
            return Response(
                content=f"Invalid model_json: {e}",
                media_type="text/plain",
                status_code=400,
            )
        docx_bytes = apply_model_and_generate_docx(docx_bytes, model)

    try:
        pdf_bytes = convert_docx_to_pdf_bytes(docx_bytes)
    except Exception as e:
        return Response(
            content=f"Preview conversion error: {e}",
            media_type="text/plain",
            status_code=500,
        )

    return Response(content=pdf_bytes, media_type="application/pdf")


from fastapi.staticfiles import StaticFiles

# Mount the static folder so:
#   GET /              -> serves static/index.html
#   GET /reset-password.html -> serves that file
app.mount("/", StaticFiles(directory="static", html=True), name="static")
