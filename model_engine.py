from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
from copy import deepcopy
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run as DocxRun


# ===================== BASIC DOCX HELPERS =====================

def load_document_from_bytes(data: bytes) -> Document:
    """Create a python-docx Document from raw .docx bytes."""
    file_obj = BytesIO(data)
    return Document(file_obj)


def _is_heading_text(text: str, keyword: str) -> bool:
    """
    True if paragraph text should be treated as a heading
    for the given keyword.

    - Case-insensitive
    - Strips spaces and trailing ':' (e.g., "PROFESSIONAL EXPERIENCE:")
    """
    if not text:
        return False
    t = text.strip().upper()
    k = keyword.strip().upper()
    t_stripped = t.rstrip(" :")
    return t_stripped == k


def find_heading_index(paragraphs: List[DocxParagraph], keywords: List[str]) -> Optional[int]:
    """Find index of first paragraph that matches any heading keyword."""
    for idx, p in enumerate(paragraphs):
        text = p.text
        if not text or not text.strip():
            continue
        for key in keywords:
            if _is_heading_text(text, key):
                return idx
    return None


def copy_run_format(src: Optional[DocxRun], dst: DocxRun) -> None:
    """
    Copy only basic emphasis flags from src run to dst run.

    IMPORTANT:
    - Do NOT copy font name, size or color.
      We let the original paragraph style control those so the
      resume keeps its original fonts/colours (blue name, etc.).
    """
    if src is None or dst is None:
        return
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline


def replace_paragraph_text_keep_style(para: DocxParagraph, new_text: str) -> None:
    """
    Change the visible text of a paragraph but keep its style,
    bullets and run formatting (bold/italic/underline) at the paragraph level.

    We only overwrite the first run's text, then clear others.
    """
    if not para.runs:
        run = para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def set_paragraph_text_with_bold_markers(para: DocxParagraph, text: str) -> None:
    """
    Replace a paragraph's runs based on **bold** markers in the text.

    Example:
      'Built **REST APIs** using **.NET 8** and **SQL Server 2022**.'
    becomes multiple runs where segments between ** ** are bold.

    We preserve formatting from the original first run and only toggle bold.
    """
    template_run = para.runs[0] if para.runs else None

    # remove existing runs
    for r in list(para.runs):
        r._element.getparent().remove(r._element)

    if not text:
        return

    parts = text.split("**")
    # Even index -> normal text, odd index -> bold text
    for i, segment in enumerate(parts):
        if not segment:
            continue
        run = para.add_run(segment)
        if template_run is not None:
            copy_run_format(template_run, run)
        run.bold = bool(i % 2 == 1)


def _run_has_line_break(run: DocxRun) -> bool:
    """
    True if this run contains a line break (<w:br/>).
    Used to detect logical lines inside a single paragraph header (Case 1).
    """
    for child in run._r:
        # child.tag looks like '{namespace}br'
        if child.tag.endswith("br"):
            return True
    return False


def _update_single_paragraph_header_role(para: DocxParagraph, new_role: str) -> None:
    """
    Case 1 (claims-style) header:

      One paragraph with 3 logical lines separated by <w:br/>:
        1) Name
        2) Role
        3) Contact line

    We:
      - Walk the paragraph XML children,
      - Count <w:br/> nodes,
      - Treat all <w:t> text nodes between the 1st and 2nd <w:br/> as the role line,
      - Replace their text with new_role (first node gets full text, others cleared),
      - Leave name + contact (and hyperlinks) untouched.
    """
    if not isinstance(new_role, str):
        return
    new_role = new_role.strip()
    if not new_role:
        return

    # We need at least two line breaks in the paragraph to have a distinct role line.
    # We'll collect <w:t> nodes that belong to the text between the 1st and 2nd <w:br/>.
    br_seen = 0
    in_role_segment = False
    role_text_elems = []

    for run in para.runs:
        r = run._r
        for child in r:
            # child.tag is like '{namespace}t' or '{namespace}br'
            local = child.tag.rsplit("}", 1)[-1]

            if local in ("br", "cr"):
                br_seen += 1
                if br_seen == 1:
                    # Just passed the first line break → now inside the role line
                    in_role_segment = True
                elif br_seen == 2:
                    # Just passed the second line break → role line is finished
                    in_role_segment = False

            elif local == "t" and in_role_segment:
                role_text_elems.append(child)

    if not role_text_elems:
        # Could not locate distinct role line safely – do nothing.
        return

    # Put the new role text into the first text node in that segment, blank out any others.
    first = True
    for t_elem in role_text_elems:
        if first:
            t_elem.text = new_role
            first = False
        else:
            t_elem.text = ""


def _parse_label_value_line(text: str, default_label: str) -> Dict[str, str]:
    """
    For lines like:
      'Environment: .NET 7/6, C#, SQL Server'
      'Role: Senior AWS Cloud/Network Engineer'
      'Project Description: Deployment of ...'

    Returns a dict: {"label": "Environment:", "value": "..."}
    """
    if ":" in text:
        label, value = text.split(":", 1)
        label = label.strip()
        value = value.strip()
    else:
        label = default_label.strip()
        value = text.strip()
    if not label.endswith(":"):
        label += ":"
    return {"label": label, "value": value}


def _looks_like_role_title(text: str) -> bool:
    """
    Heuristic to detect a 'role' line when there is no explicit 'Role:' label.
    Example: '.NET Developer', 'Senior Java Full Stack Developer'.

    Rules:
      - No colon
      - Not too long (<= 7 words)
      - Contains at least one common role keyword
    """
    if not text:
        return False

    s = text.strip()
    if not s:
        return False

    # Skip lines that already look like 'Label: Value'
    if ":" in s:
        return False

    words = s.split()
    if len(words) == 0 or len(words) > 7:
        return False

    lowered = s.lower()
    role_keywords = [
        "developer",
        "engineer",
        "analyst",
        "consultant",
        "architect",
        "specialist",
        "administrator",
        "lead",
        "manager",
    ]
    return any(k in lowered for k in role_keywords)


def _smart_split_single_paragraph_header(text: str) -> List[str]:
    """
    Try to split a single big header line into:
        [name, title, contact...]
    for Case 1 style headers like:
        "Abdul Sattar Shaik EDI Specialist +1 (848) 300-5996 | email | LinkedIn"

    If we can't confidently split, return [] so caller can keep original.
    """
    s = (text or "").strip()
    if not s:
        return []

    # 1) If explicit newlines exist, prefer them
    if "\n" in s:
        parts = [p.strip() for p in s.split("\n") if p.strip()]
        if len(parts) >= 2:
            return parts

    words = s.split()
    if not words:
        return []

    # 2) Find where contact info starts (phone/email/LinkedIn/link-ish)
    contact_idx = None
    for i, w in enumerate(words):
        wl = w.lower()

        # email / linkedin / url
        if "@" in wl or "linkedin" in wl or "http" in wl or "www." in wl:
            contact_idx = i
            break

        # phone-like: token with multiple digits or explicit '+'-phone
        digits = sum(ch.isdigit() for ch in wl)
        if wl.startswith("+") and digits >= 1:
            contact_idx = i
            break
        if digits >= 3:  # "(848)", "300-5996", etc.
            contact_idx = i
            break

    if contact_idx is None:
        # No obvious contact marker – treat entire string as "name + title"
        contact_idx = len(words)

    prefix_words = words[:contact_idx]
    contact_words = words[contact_idx:]

    # 3) Inside the prefix, find where the role/title likely starts
    role_keywords = [
        "developer",
        "engineer",
        "analyst",
        "consultant",
        "architect",
        "specialist",
        "administrator",
        "lead",
        "manager",
        "scientist",
        "edi",
        ".net",
        "java",
        "full",
        "stack",
        "cloud",
        "data",
    ]

    role_start = None
    for i, w in enumerate(prefix_words):
        wl = re.sub(r"[^\w]+", "", w).lower()
        if any(k in wl for k in role_keywords):
            role_start = i
            break

    if role_start is None:
        # Fallback: assume first 2 words are name, rest is title (if any)
        if len(prefix_words) <= 2:
            name_words = prefix_words
            title_words = []
        else:
            name_words = prefix_words[:2]
            title_words = prefix_words[2:]
    else:
        name_words = prefix_words[:role_start]
        title_words = prefix_words[role_start:]

    line_name = " ".join(name_words).strip()
    line_title = " ".join(title_words).strip()
    line_contact = " ".join(contact_words).strip()

    result: List[str] = []
    if line_name:
        result.append(line_name)
    if line_title:
        result.append(line_title)
    if line_contact:
        result.append(line_contact)

    # Only use this if we got at least 2 logical pieces
    if len(result) >= 2:
        return result

    # Not confident → let caller keep original
    return []


# ===================== EXPERIENCE HELPERS =====================

def is_probable_job_header(para: DocxParagraph) -> bool:
    """
    Heuristic to detect a job header line inside the EXPERIENCE section.

    We treat a paragraph as a job header if:
      - It has text.
      - It is *not* a bullet/numbered list.
      - A clear majority of its runs are bold.
      - It contains at least one digit (dates).
    """
    text = (para.text or "").strip()
    if not text:
        return False

    # Require at least one digit (dates) to consider it a header
    if not any(ch.isdigit() for ch in text):
        return False

    # 1) Do NOT treat bullets / list paragraphs as headers
    style_name = (para.style.name or "").lower()
    if "list" in style_name or "bullet" in style_name:
        return False

    # 2) Check for numbering/bullets via underlying XML (numPr)
    pPr = getattr(para._p, "pPr", None)
    if pPr is not None and getattr(pPr, "numPr", None) is not None:
        return False

    runs = [r for r in para.runs if (r.text or "").strip()]
    if not runs:
        return False

    bold_count = sum(1 for r in runs if r.bold)
    return bold_count > 0 and bold_count >= len(runs) * 0.6


def is_probable_role_line(para: DocxParagraph, text: Optional[str] = None) -> bool:
    """
    Detect a 'Role line' such as 'Sr .NET Full Stack Developer' that:
      - Is bold
      - Is not a list/bullet
      - Contains NO digits
      - Does not start with 'Client:', 'Environment:', etc.
    """
    if text is None:
        text = para.text or ""
    stripped = text.strip()
    if not stripped:
        return False

    upper = stripped.upper()

    # Exclude obvious labels that we handle elsewhere
    if upper.startswith("CLIENT:"):
        return False
    if upper.startswith("ENVIRONMENT:"):
        return False
    if upper.startswith("PROJECT DESCRIPTION:"):
        return False
    if upper.startswith("RESPONSIBILITIES:"):
        return False
    if upper.startswith("ROLE:"):
        return False

    # Require: no digits
    if any(ch.isdigit() for ch in stripped):
        return False

    # Exclude section headings
    if len(stripped) <= 3:
        return False

    # Exclude bullets
    style_name = (para.style.name or "").lower()
    if "list" in style_name or "bullet" in style_name:
        return False
    pPr = getattr(para._p, "pPr", None)
    if pPr is not None and getattr(pPr, "numPr", None) is not None:
        return False

    runs = [r for r in para.runs if (r.text or "").strip()]
    if not runs:
        return False

    bold_count = sum(1 for r in runs if r.bold)
    return bold_count > 0 and bold_count >= len(runs) * 0.6


SECTION_HEADING_WORDS = {
    "SUMMARY",
    "PROFESSIONAL SUMMARY",
    "SKILLS",
    "TECHNICAL SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "EXPERIENCE",
    "WORK EXPERIENCE",
    "EMPLOYMENT HISTORY",
    "EDUCATION",
    "EDUCATION DETAILS",
    "CERTIFICATIONS",
    "CERTIFICATION",
    "PROJECTS",
    "PROJECT EXPERIENCE",
    "ACADEMIC PROJECTS",
}


def _is_heading_like(text: str) -> bool:
    """
    Returns True if a line looks like a top-level section heading (all caps / short).

    We use this both when cleaning experience bullets and when scanning
    the whole document for generic sections.
    """
    if not text:
        return False
    stripped = text.strip()
    upper = stripped.upper()

    # Reject lines with digits – section headings rarely contain dates/numbers.
    # This helps avoid treating job header lines with dates as top-level headings.
    if any(ch.isdigit() for ch in stripped):
        return False

    # Explicit known headings
    if upper in SECTION_HEADING_WORDS:
        return True

    # All caps and not too long – e.g. "WORK EXPERIENCE" / "WHO THE HELL"
    if upper == stripped and len(stripped.split()) <= 6:
        return True

    return False


def _is_responsibilities_label(text: str) -> bool:
    """
    Detect 'Responsibilities:' label lines (with or without bullet chars),
    even with minor typos. These should stay as bold, non-bullet lines.
    """
    if not text:
        return False

    t = text.strip()
    # Strip common bullet characters
    t = t.lstrip("•*-\u2022·").strip()

    # Drop trailing colon or period
    t = t.rstrip(":").rstrip(".").strip()
    if not t:
        return False

    lt = t.lower()

    # Exact matches
    if lt in ("responsibilities", "key responsibilities"):
        return True

    # Allow small spelling variations
    if lt.startswith("responsibil"):
        return True
    if "responsibilit" in lt:
        return True

    return False


def _next_heading_for_other_sections(paragraphs: List[DocxParagraph], start_idx: int) -> Optional[int]:
    """
    For OTHER / generic sections, find the next *top-level* heading after start_idx.
    We skip ENVIRONMENT:, ROLE:, PROJECT DESCRIPTION:, and Responsibilities-type labels
    because those live inside EXPERIENCE job blocks.
    """
    for idx in range(start_idx + 1, len(paragraphs)):
        txt = (paragraphs[idx].text or "").strip()
        if not txt:
            continue

        if _is_responsibilities_label(txt):
            continue

        upper = txt.upper()
        if upper.startswith("ENVIRONMENT:") or upper.startswith("ROLE:") or upper.startswith("PROJECT DESCRIPTION:"):
            continue

        if _is_heading_like(txt):
            return idx

    return None


def clean_experience_model(model: Dict[str, Any]) -> None:
    """
    In-place cleanup of the experience section in the JSON model.
    - Removes bullets that are actually headings like 'EDUCATION DETAILS'.
    - Pulls Environment / Role / Project Description lines into objects.
    """
    if "experience" not in model:
        return

    new_experience: List[Dict[str, Any]] = []

    for job in model["experience"]:
        header_text = (job.get("header") or "").strip()
        bullets = job.get("bullets", []) or []
        environment = job.get("environment")
        role = job.get("role")
        proj = job.get("project_description")

        cleaned_bullets: List[str] = []

        for b in bullets:
            if not b:
                continue
            t = b.strip()
            upper = t.upper()

            # drop obvious headings (EDUCATION, CERTIFICATIONS, etc.)
            if _is_heading_like(t):
                continue

            # Fallback extraction for environment/role/project_description
            if upper.startswith("ENVIRONMENT:"):
                if not environment:
                    environment = _parse_label_value_line(t, "Environment")
                continue
            if upper.startswith("ROLE:"):
                if not role:
                    role = _parse_label_value_line(t, "Role")
                continue
            if upper.startswith("PROJECT DESCRIPTION:"):
                if not proj:
                    proj = _parse_label_value_line(t, "Project Description")
                continue

            cleaned_bullets.append(b)

        # --- Fallback 1: Role from bullets when role is a bullet ---
        if not role:
            fallback_role_text: Optional[str] = None
            remaining_bullets: List[str] = []

            for t in cleaned_bullets:
                if fallback_role_text is None and _looks_like_role_title(t):
                    fallback_role_text = t.strip()
                    continue
                remaining_bullets.append(t)

            cleaned_bullets = remaining_bullets
            if fallback_role_text:
                role = {
                    "label": "Role:",
                    "value": fallback_role_text,
                }

        # --- Fallback 2: Role from header second line ---
        if not role and header_text:
            header_lines = header_text.splitlines()
            if len(header_lines) >= 2:
                last_line = header_lines[-1].strip()
                if _looks_like_role_title(last_line):
                    role = {
                        "label": "Role:",
                        "value": last_line,
                    }
                    # Keep the rest as the header (company + location + dates)
                    new_header = "\n".join(header_lines[:-1]).rstrip()
                    if new_header:
                        job["header"] = new_header
                    else:
                        job["header"] = header_text

        job["bullets"] = cleaned_bullets
        if environment:
            job["environment"] = environment
        if role:
            job["role"] = role
        if proj:
            job["project_description"] = proj

        new_experience.append(job)

    model["experience"] = new_experience


def _infer_resume_template_from_text(text: str) -> str:
    """
    Very simple heuristic classifier for high-level resume template:
    - case1_edi_claims      -> EDI / HIPAA / 837 claims-heavy
    - case2_dotnet          -> .NET / C# / Web API banking-ish
    - case3_networking      -> networking/cloud with routing protocols
    - generic               -> fallback
    """
    if not isinstance(text, str):
        return "generic"

    low = text.lower()

    # Case 1: EDI / X12 / HIPAA claims world
    if "x12" in low or "hipaa" in low or "837p" in low or "837i" in low or "edi 837" in low:
        return "case1_edi_claims"

    # Case 2: .NET / C# / Web API / banking style
    if ".net" in low or "asp.net" in low or "c#" in low or "web api" in low:
        return "case2_dotnet"

    # Case 3: Networking / routing / cisco
    if "bgp" in low or "ospf" in low or "eigrp" in low or "cisco" in low:
        return "case3_networking"

    return "generic"


def _table_looks_like_skills(table) -> bool:
    """
    Heuristic: return True if this table is likely a skills table,
    not a project/experience table.

    We avoid mis-tagging project/experience tables as skills by:
      - Rejecting tables that talk a lot about client/project/company/responsibilities.
      - Prefer tables that mention skills/technical/tools/etc.
      - As a fallback, treat tables with many short, comma-heavy cells as skills-ish.
    """
    texts: List[str] = []
    for row in table.rows:
        for cell in row.cells:
            t = (cell.text or "").strip()
            if t:
                texts.append(t)

    if not texts:
        return False

    joined = " ".join(texts).lower()

    # Strong signals that this is NOT a skills table
    bad_markers = ["client:", "project:", "company:", "responsibilities:"]
    if any(m in joined for m in bad_markers):
        return False

    # Strong signals that this IS a skills table
    skill_markers = [
        "skills",
        "technical",
        "tools",
        "technologies",
        "frameworks",
        "languages",
        "databases",
        "cloud",
    ]
    if any(m in joined for m in skill_markers):
        return True

    # Fallback: lots of short, comma-heavy cells → skills list vibes
    short_cells = [t for t in texts if len(t) <= 60]
    if not short_cells:
        return False

    comma_count = sum(t.count(",") for t in short_cells)
    comma_ratio = comma_count / max(len(short_cells), 1)

    if comma_ratio >= 0.5:
        return True

    return False


# ===================== PARSE DOCX -> MODEL =====================

def parse_docx_to_model(doc: Document) -> Dict[str, Any]:
    """
    Build a simple model from the DOCX that we can edit via JSON.

    Sections we support:
      - HEADER (all non-empty lines before the first heading)
      - SUMMARY bullets
      - SKILLS section:
          * Claims-style: table with skills
          * Paragraph-style: block after SKILLS / TECHNICAL SKILLS
      - EXPERIENCE: list of blocks with:
          * header
          * bullets
          * optional environment (label/value)
          * optional role (label/value)
          * optional project_description (label/value)
      - EDUCATION / CERTIFICATIONS / PROJECTS as simple lines
      - OTHER SECTIONS: any unknown heading blocks (AWARDS, DETAILS, WHO THE HELL ARE YOU, etc.)
    """
    model: Dict[str, Any] = {}
    paragraphs = list(doc.paragraphs)

    # ---------- FIND SECTION HEADINGS (core known sections) ----------
    summary_idx = find_heading_index(paragraphs, ["SUMMARY", "PROFESSIONAL SUMMARY"])
    skills_heading_idx = find_heading_index(paragraphs, ["SKILLS", "TECHNICAL SKILLS"])
    experience_idx = find_heading_index(
        paragraphs,
        ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
    )
    education_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])
    cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
    project_idx = find_heading_index(
        paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"]
    )

    # ---------- GLOBAL HEADING MAP (known + unknown) ----------
    # Build a sorted list of all heading-like paragraphs.
    # Entry: (paragraph_index, text, is_core_section, core_section_key_or_None)
    all_heading_entries: List[Tuple[int, str, bool, Optional[str]]] = []

    known_heading_map: Dict[int, str] = {}
    if summary_idx is not None:
        known_heading_map[summary_idx] = "summary"
    if skills_heading_idx is not None:
        known_heading_map[skills_heading_idx] = "skills"
    if experience_idx is not None:
        known_heading_map[experience_idx] = "experience"
    if education_idx is not None:
        known_heading_map[education_idx] = "education"
    if cert_idx is not None:
        known_heading_map[cert_idx] = "certifications"
    if project_idx is not None:
        known_heading_map[project_idx] = "projects"

    for i, p in enumerate(paragraphs):
        txt = (p.text or "").strip()
        if not txt:
            continue

        # Skip responsibilities / env / role / project-description labels;
        # these live inside jobs, not as top-level resume sections.
        if _is_responsibilities_label(txt):
            continue
        upper = txt.upper()
        if upper.startswith("ENVIRONMENT:") or upper.startswith("ROLE:") or upper.startswith(
            "PROJECT DESCRIPTION:"
        ):
            continue

        # Always include core headings if we know their indices
        if i in known_heading_map:
            all_heading_entries.append((i, txt, True, known_heading_map[i]))
            continue

        if not _is_heading_like(txt):
            continue

        # This is a generic / unknown heading (e.g. AWARDS, DETAILS, WHO THE HELL ARE YOU)
        all_heading_entries.append((i, txt, False, None))

    all_heading_entries.sort(key=lambda x: x[0])

    def _next_heading_after(idx: int) -> Optional[int]:
        """
        Return the paragraph index of the next heading (known or unknown)
        strictly after idx, or None if there is no later heading.
        """
        for hi, _, _, _ in all_heading_entries:
            if hi > idx:
                return hi
        return None

    # ---------- UNIFIED HEADING MAP FOR CORE SECTIONS (ANY ORDER) ----------
    # Used by get_section_bounds() so sections can appear in any order.
    section_positions: List[Tuple[str, int]] = []
    for key, idx in [
        ("summary", summary_idx),
        ("skills", skills_heading_idx),
        ("experience", experience_idx),
        ("education", education_idx),
        ("certifications", cert_idx),
        ("projects", project_idx),
    ]:
        if idx is not None:
            section_positions.append((key, idx))

    section_positions.sort(key=lambda x: x[1])

    def get_section_bounds(section_key: str) -> Tuple[Optional[int], Optional[int]]:
        """
        Return (start_idx, end_idx) for a logical core section:

          - start_idx: first paragraph AFTER the section heading.
          - end_idx:   index of the NEXT heading (known or unknown),
                      or len(paragraphs) if this heading is the last one.

        If the section heading does not exist, returns (None, None).
        """
        for key, idx in section_positions:
            if key == section_key:
                start = idx + 1
                next_idx_any = _next_heading_after(idx)
                end = next_idx_any if next_idx_any is not None else len(paragraphs)
                return start, end
        return None, None

    # ---------- HEADER (all non-empty lines before the first heading) ----------
    header_lines: List[str] = []
    header_indices: List[int] = []

    if all_heading_entries:
        # First heading (core OR generic, e.g. SUMMARY, DETAILS, AWARDS, WHO THE HELL ARE YOU, ...)
        header_end = all_heading_entries[0][0]
    else:
        # No headings at all – header will be empty.
        header_end = 0

    max_header_paras = 3  # HARD CAP: at most 3 header paragraphs

    for idx in range(0, header_end):
        p = paragraphs[idx]
        txt = (p.text or "").strip()
        if not txt:
            continue

        # Do NOT allow bullets / numbered list items into the header.
        style_name = (p.style.name or "").lower()
        pPr = getattr(p._p, "pPr", None)
        is_list = ("list" in style_name) or ("bullet" in style_name)
        if pPr is not None and getattr(pPr, "numPr", None) is not None:
            is_list = True

        if is_list:
            # First bullet we see → stop header here.
            break

        header_lines.append(txt)
        header_indices.append(idx)

        if len(header_indices) >= max_header_paras:
            break

    if header_lines:
        # Determine header layout:
        #   - single_paragraph: one paragraph (Case 1 / claims header)
        #   - multi_paragraph:  multiple lines (Case 2/3)
        if len(header_indices) == 1:
            header_layout = "single_paragraph"

            # For Case 1, try to split single big line into [Name, Title, Contact]
            original = header_lines[0]
            split_lines = _smart_split_single_paragraph_header(original)
            if split_lines:
                header_lines = split_lines
        else:
            header_layout = "multi_paragraph"

        model["header"] = {"lines": header_lines}
        model["header_meta"] = {
            "indices": header_indices,
            "layout": header_layout,
        }

    # ---------- SUMMARY ----------
    summary_bullets: List[str] = []
    summary_indices: List[int] = []

    # ONLY if there is an explicit SUMMARY heading
    if summary_idx is not None:
        start_idx, end_idx = get_section_bounds("summary")
        if start_idx is not None and end_idx is not None:
            for idx in range(start_idx, end_idx):
                text = (paragraphs[idx].text or "").strip()
                if text:
                    summary_bullets.append(text)
                    summary_indices.append(idx)

    if summary_bullets:
        model["summary"] = {"bullets": summary_bullets}
        model["summary_meta"] = {"indices": summary_indices}

    # ---------- SKILLS ----------
    skills_model: Dict[str, Any] = {"columns": 2, "rows": []}
    skills_meta: Dict[str, Any] = {}

    if skills_heading_idx is not None:
        lines: List[str] = []
        start_idx, end_idx = get_section_bounds("skills")
        if start_idx is not None and end_idx is not None:
            for idx in range(start_idx, end_idx):
                text = paragraphs[idx].text.strip()
                if text:
                    lines.append(text)

        if lines:
            # Paragraph-style skills + multi-label block detection
            rows_data: List[List[str]] = []
            multi_label_rows: List[List[str]] = []
            multi_label_detected = False

            # Some resumes put ALL skills lines into a single paragraph with soft line breaks.
            label_pattern = r"[A-Z][A-Za-z0-9 &/()+]*:"

            for ln in lines:
                ln_strip = ln.strip()
                if not ln_strip:
                    continue

                matches = list(re.finditer(label_pattern, ln_strip))
                if len(matches) >= 2:
                    # Treat this as a multi-label block; each label:value becomes its own row.
                    multi_label_detected = True
                    for i, m in enumerate(matches):
                        start = m.start()
                        end = matches[i + 1].start() if i + 1 < len(matches) else len(ln_strip)
                        part = ln_strip[start:end].strip()
                        if ":" not in part:
                            continue
                        left, right = part.split(":", 1)
                        multi_label_rows.append([left.strip(), right.strip()])
                else:
                    # Normal single-label line or plain text
                    if ":" in ln_strip:
                        left, right = ln_strip.split(":", 1)
                        rows_data.append([left.strip(), right.strip()])
                    else:
                        rows_data.append(["", ln_strip])

            if multi_label_detected and multi_label_rows:
                skills_model["rows"] = multi_label_rows
                skills_meta = {"format": "paragraph", "layout": "multi_label_block"}
            else:
                skills_model["rows"] = rows_data
                skills_meta = {"format": "paragraph"}

        elif doc.tables:
            # Heading present but no text → fallback to first table
            table = doc.tables[0]
            rows_data: List[List[str]] = []
            for row in table.rows:
                if len(row.cells) >= 2:
                    left = row.cells[0].text.strip()
                    right = row.cells[1].text.strip()
                elif len(row.cells) == 1:
                    left = row.cells[0].text.strip()
                    right = ""
                else:
                    left = right = ""
                rows_data.append([left, right])
            skills_model["rows"] = rows_data
            skills_meta = {"format": "table", "table_index": 0}

    # No SKILLS heading at all → try to detect a skills *table* heuristically
    elif doc.tables:
        chosen_index: Optional[int] = None

        for idx, table in enumerate(doc.tables):
            if _table_looks_like_skills(table):
                chosen_index = idx
                break

        if chosen_index is not None:
            table = doc.tables[chosen_index]
            rows_data: List[List[str]] = []
            for row in table.rows:
                if len(row.cells) >= 2:
                    left = row.cells[0].text.strip()
                    right = row.cells[1].text.strip()
                elif len(row.cells) == 1:
                    left = row.cells[0].text.strip()
                    right = ""
                else:
                    left = right = ""
                rows_data.append([left, right])

            skills_model["rows"] = rows_data
            skills_meta = {"format": "table", "table_index": chosen_index}

    if skills_model["rows"]:
        model["skills"] = skills_model
        model["skills_meta"] = skills_meta

    # ---------- EXPERIENCE ----------
    experience_blocks: List[Dict[str, Any]] = []

    if experience_idx is not None:
        start_idx, exp_end_idx = get_section_bounds("experience")
        if start_idx is None or exp_end_idx is None:
            start_idx = experience_idx + 1
            exp_end_idx = len(paragraphs)

        current_header: Optional[str] = None
        current_bullets: List[str] = []
        current_env: Optional[Dict[str, str]] = None
        current_role: Optional[Dict[str, str]] = None
        current_proj_desc: Optional[Dict[str, str]] = None

        for idx in range(start_idx, exp_end_idx):
            para = paragraphs[idx]
            text = para.text.strip()
            if not text:
                continue

            upper = text.upper()
            stripped_upper = upper.lstrip("•*-\u2022· ").strip()
            words = stripped_upper.split()
            first_word = words[0].rstrip(":") if words else ""

            # --- New job header? ---
            if is_probable_job_header(para):
                # Skip "Role:" and "Project Description:" as new headers
                if first_word not in {"ROLE", "PROJECT"} and not stripped_upper.startswith("ENVIRONMENT:"):
                    if current_header is not None:
                        job: Dict[str, Any] = {
                            "header": current_header,
                            "bullets": current_bullets,
                        }
                        if current_env:
                            job["environment"] = current_env
                        if current_role:
                            job["role"] = current_role
                        if current_proj_desc:
                            job["project_description"] = current_proj_desc
                        experience_blocks.append(job)

                    current_header = text
                    current_bullets = []
                    current_env = None
                    current_role = None
                    current_proj_desc = None
                    continue

            # --- Inside a job block ---
            if current_header is not None:
                # Case 2 (.NET): bold role line immediately after header
                if current_role is None and is_probable_role_line(para, text):
                    current_role = {
                        "label": "Role:",
                        "value": text.strip(),
                    }
                    continue

                if stripped_upper.startswith("ENVIRONMENT:"):
                    current_env = _parse_label_value_line(text, "Environment")
                    continue

                if stripped_upper.startswith("ROLE:"):
                    current_role = _parse_label_value_line(text, "Role")
                    continue

                if stripped_upper.startswith("PROJECT DESCRIPTION:"):
                    parsed = _parse_label_value_line(text, "Project Description")
                    if current_proj_desc is None:
                        current_proj_desc = parsed
                    else:
                        val = parsed.get("value", "")
                        if val:
                            current_proj_desc["value"] = (
                                (current_proj_desc["value"] + " " + val).strip()
                            )
                    continue

                # Normal bullet/content (including 'Responsibilities:' labels for now)
                current_bullets.append(text)

        # Final job block
        if current_header is not None:
            job = {
                "header": current_header,
                "bullets": current_bullets,
            }
            if current_env:
                job["environment"] = current_env
            if current_role:
                job["role"] = current_role
            if current_proj_desc:
                job["project_description"] = current_proj_desc
            experience_blocks.append(job)

    if experience_blocks:
        model["experience"] = experience_blocks

    # Final cleanup on experience
    clean_experience_model(model)

    # ---------- EXPERIENCE META ----------
    if "experience" in model:
        has_env = False
        has_role = False
        has_proj_desc = False

        for job in model["experience"]:
            if not isinstance(job, dict):
                continue

            env = job.get("environment") or {}
            role = job.get("role") or {}
            proj = job.get("project_description") or {}

            if isinstance(env, dict) and env.get("value"):
                has_env = True
            if isinstance(role, dict) and role.get("value"):
                has_role = True
            if isinstance(proj, dict) and proj.get("value"):
                has_proj_desc = True

        model["experience_meta"] = {
            "has_environment_label": has_env,
            "has_role_label": has_role,
            "has_project_description_label": has_proj_desc,
        }

    # ---------- EDUCATION ----------
    if education_idx is not None:
        edu_lines: List[str] = []
        start_idx, edu_end_idx = get_section_bounds("education")
        if start_idx is not None and edu_end_idx is not None:
            for idx in range(start_idx, edu_end_idx):
                text = paragraphs[idx].text.strip()
                if text:
                    edu_lines.append(text)

        if edu_lines:
            model["education"] = {"lines": edu_lines}

    # ---------- CERTIFICATIONS ----------
    if cert_idx is not None:
        cert_lines: List[str] = []
        start_idx, cert_end_idx = get_section_bounds("certifications")
        if start_idx is not None and cert_end_idx is not None:
            for idx in range(start_idx, cert_end_idx):
                text = paragraphs[idx].text.strip()
                if text:
                    cert_lines.append(text)

        if cert_lines:
            model["certifications"] = {"lines": cert_lines}

    # ---------- PROJECTS ----------
    if project_idx is not None:
        proj_lines: List[str] = []
        start_idx, proj_end_idx = get_section_bounds("projects")
        if start_idx is not None and proj_end_idx is not None:
            for idx in range(start_idx, proj_end_idx):
                text = paragraphs[idx].text.strip()
                if text:
                    proj_lines.append(text)

        if proj_lines:
            model["projects"] = {"lines": proj_lines}

    # ---------- OTHER / GENERIC SECTIONS (non-core headings) ----------
    # Any heading-like block that is NOT one of:
    #   summary / skills / experience / education / certifications / projects
    # becomes an "other section" we can round-trip.
    other_sections: List[Dict[str, Any]] = []

    if all_heading_entries:
        for hi, title, is_known, section_key in all_heading_entries:
            if is_known:
                # Skip core sections; they are already modeled separately.
                continue

            # Content for this section = paragraphs between this heading
            # and the next heading (of any kind).
            start = hi + 1
            next_idx = _next_heading_after(hi)
            end = next_idx if next_idx is not None else len(paragraphs)

            lines: List[str] = []
            for i in range(start, end):
                txt = (paragraphs[i].text or "").strip()
                if txt:
                    lines.append(txt)

            if not lines:
                continue

            # ID is a slug version of the heading
            sec_id = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_") or f"section_{hi}"

            other_sections.append(
                {
                    "id": sec_id,
                    "title": title.strip(),
                    "kind": "lines",
                    "lines": lines,
                    "meta": {
                        "heading_original_text": title.strip(),
                    },
                }
            )

    if other_sections:
        model["other_sections"] = other_sections

    # ---------- TEMPLATE META (layout + rough template guess) ----------
    try:
        template_meta: Dict[str, Any] = {}

        # Pull meta we already computed
        header_meta = model.get("header_meta") or {}
        skills_meta = model.get("skills_meta") or {}
        exp_meta = model.get("experience_meta") or {}

        if not isinstance(header_meta, dict):
            header_meta = {}
        if not isinstance(skills_meta, dict):
            skills_meta = {}
        if not isinstance(exp_meta, dict):
            exp_meta = {}

        header_layout = header_meta.get("layout")
        header_indices_meta = header_meta.get("indices") or []

        skills_format = skills_meta.get("format")
        skills_layout = skills_meta.get("layout")
        skills_table_index = skills_meta.get("table_index")

        has_env = bool(exp_meta.get("has_environment_label"))
        has_role = bool(exp_meta.get("has_role_label"))
        has_proj_desc = bool(exp_meta.get("has_project_description_label"))

        # Section order based on heading indices we already computed
        section_indices: List[Tuple[str, int]] = []
        for name, idx in [
            ("summary", summary_idx),
            ("skills", skills_heading_idx),
            ("experience", experience_idx),
            ("education", education_idx),
            ("certifications", cert_idx),
            ("projects", project_idx),
        ]:
            if idx is not None:
                section_indices.append((name, idx))

        section_indices.sort(key=lambda x: x[1])
        section_order = [name for name, _ in section_indices]

        # Whole-text sniffing for template case
        try:
            full_text_parts: List[str] = []
            for p in doc.paragraphs:
                if p.text:
                    full_text_parts.append(p.text)
            full_text = "\n".join(full_text_parts)
        except Exception:
            full_text = ""

        resume_template = _infer_resume_template_from_text(full_text)

        template_meta.update(
            {
                "resume_template": resume_template,
                "header_layout": header_layout,
                "header_indices": header_indices_meta,
                "skills_format": skills_format,
                "skills_layout": skills_layout,
                "skills_table_index": skills_table_index,
                "has_environment_label": has_env,
                "has_role_label": has_role,
                "has_project_description_label": has_proj_desc,
                "section_order": section_order,
            }
        )

        model["template_meta"] = template_meta
        # Also store a top-level section_order so apply_model_to_docx can use it
        model["section_order"] = section_order

    except Exception:
        # Template detection should never break parsing; fail silently.
        pass

    return model


# ===================== APPLY MODEL -> DOCX =====================

def update_bullet_paragraphs(bullet_paras: List[DocxParagraph], new_bullets: List[str]) -> None:
    """
    Shared logic for updating bullet/line paragraphs:

      - We match by TEXT, not just by index.
      - Remove original paragraphs and reinsert or clone them with new text.
      - Supports **bold** markers in the new text.
    """
    if not bullet_paras and not new_bullets:
        return

    if not bullet_paras:
        return

    def norm(text: str) -> str:
        return (text or "").strip()

    # Map text -> list of paragraphs
    text_to_paras: Dict[str, List[DocxParagraph]] = {}
    for p in bullet_paras:
        t = norm(p.text)
        text_to_paras.setdefault(t, []).append(p)

    # Find parent and anchor
    first_elem = bullet_paras[0]._element
    parent = first_elem.getparent()
    if parent is None:
        return

    anchor_before = first_elem.getprevious()
    template_el = deepcopy(first_elem)

    # Remove all original paragraphs from document
    for p in bullet_paras:
        elem = p._element
        if elem.getparent() is parent:
            parent.remove(elem)

    last_inserted = anchor_before

    def insert_after(elem_to_insert, after_elem):
        if after_elem is None:
            parent.insert(0, elem_to_insert)
        else:
            parent.insert(parent.index(after_elem) + 1, elem_to_insert)

    for new_text in new_bullets:
        t_new = norm(new_text)
        reused_para: Optional[DocxParagraph] = None

        # Try reuse
        if t_new in text_to_paras and text_to_paras[t_new]:
            reused_para = text_to_paras[t_new].pop(0)

        if reused_para is not None:
            # reused paragraph: keep original runs/formatting
            elem = reused_para._element
            insert_after(elem, last_inserted)
            last_inserted = elem
        else:
            # brand new paragraph, built from template, but text may contain **bold** markers
            new_el = deepcopy(template_el)
            insert_after(new_el, last_inserted)
            last_inserted = new_el
            new_para = DocxParagraph(new_el, bullet_paras[0]._parent)
            set_paragraph_text_with_bold_markers(new_para, t_new)


def apply_summary_model(doc: Document, summary_model: Dict[str, Any], summary_meta: Dict[str, Any]) -> None:
    """
    Rewrite SUMMARY bullets only within the original bullet paragraphs.
    """
    bullets: List[str] = summary_model.get("bullets") or []
    paragraphs = list(doc.paragraphs)

    indices: List[int] = summary_meta.get("indices") or []
    bullet_paras: List[DocxParagraph] = []

    if indices:
        bullet_paras = [paragraphs[i] for i in indices if i < len(paragraphs)]
    else:
        if not bullets:
            return
        summary_idx = find_heading_index(paragraphs, ["SUMMARY", "PROFESSIONAL SUMMARY"])
        if summary_idx is None:
            return
        skills_idx = find_heading_index(paragraphs, ["SKILLS", "TECHNICAL SKILLS"])
        experience_idx = find_heading_index(
            paragraphs,
            ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
        )
        education_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])
        cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
        projects_idx = find_heading_index(
            paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "PROJECT", "PROJECT DETAILS"]
        )

        end_candidates = []
        for idx in [skills_idx, experience_idx, education_idx, cert_idx, projects_idx]:
            if idx is not None and idx > summary_idx:
                end_candidates.append(idx)
        end_idx = min(end_candidates) if end_candidates else len(paragraphs)
        bullet_paras = [p for p in paragraphs[summary_idx + 1: end_idx] if p.text.strip()]

    if not bullet_paras and not bullets:
        return

    if not bullets:
        # Remove all summary bullet paragraphs
        for p in bullet_paras:
            elem = p._element
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
        return

    if not bullet_paras:
        return

    update_bullet_paragraphs(bullet_paras, bullets)


def apply_education_model(doc: Document, edu_model: Dict[str, Any]) -> None:
    """
    Rewrite EDUCATION lines as a simple block of paragraphs.
    """
    lines: List[str] = edu_model.get("lines") or []
    if not lines:
        return

    paragraphs = list(doc.paragraphs)
    edu_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])
    if edu_idx is None:
        return

    cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
    project_idx = find_heading_index(paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"])

    end_candidates = []
    for idx in [cert_idx, project_idx]:
        if idx is not None and idx > edu_idx:
            end_candidates.append(idx)
    end_idx = min(end_candidates) if end_candidates else len(paragraphs)

    section_paras = [p for p in paragraphs[edu_idx + 1:end_idx] if p.text.strip()]
    if not section_paras:
        return

    update_bullet_paragraphs(section_paras, lines)


def apply_certifications_model(doc: Document, cert_model: Dict[str, Any]) -> None:
    """
    Rewrite CERTIFICATIONS lines as a simple block of paragraphs.
    """
    lines: List[str] = cert_model.get("lines") or []
    if not lines:
        return

    paragraphs = list(doc.paragraphs)
    cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
    if cert_idx is None:
        return

    project_idx = find_heading_index(paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"])

    end_candidates = []
    for idx in [project_idx]:
        if idx is not None and idx > cert_idx:
            end_candidates.append(idx)
    end_idx = min(end_candidates) if end_candidates else len(paragraphs)

    section_paras = [p for p in paragraphs[cert_idx + 1:end_idx] if p.text.strip()]
    if not section_paras:
        return

    update_bullet_paragraphs(section_paras, lines)


def apply_projects_model(doc: Document, proj_model: Dict[str, Any]) -> None:
    """
    Rewrite PROJECTS lines as a simple block of paragraphs.
    """
    lines: List[str] = proj_model.get("lines") or []
    if not lines:
        return

    paragraphs = list(doc.paragraphs)
    project_idx = find_heading_index(paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"])
    if project_idx is None:
        return

    section_paras = [p for p in paragraphs[project_idx + 1:] if p.text.strip()]
    if not section_paras:
        return

    update_bullet_paragraphs(section_paras, lines)


def _replace_cell_text_keep_style(cell, new_text: str) -> None:
    """
    Set the visible text of a table cell while:
      - Using the first paragraph to keep base style.
      - Clearing all other paragraphs (to avoid duplicates).
    """
    paras = list(cell.paragraphs)
    if not paras:
        p = cell.add_paragraph()
        p.add_run(new_text)
        return

    replace_paragraph_text_keep_style(paras[0], new_text)
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def apply_skills_table_model(doc: Document, skills_model: Dict[str, Any]) -> None:
    """
    Rewrite the skills table (Claims-style / table-style skills).
    """
    rows: List[List[str]] = skills_model.get("rows") or []
    if not rows or not doc.tables:
        return

    table = doc.tables[0]
    num_rows_needed = len(rows)

    # Grow table if needed
    while len(table.rows) < num_rows_needed:
        template_row = table.rows[-1]
        new_tr = deepcopy(template_row._tr)
        template_row._tr.addnext(new_tr)

    # Shrink table if needed
    while len(table.rows) > num_rows_needed:
        row = table.rows[-1]
        table._tbl.remove(row._tr)

    # Update cell contents
    for idx, (left_text, right_text) in enumerate(rows):
        row = table.rows[idx]

        left_cell = row.cells[0]
        new_left = left_text.strip()
        _replace_cell_text_keep_style(left_cell, new_left)

        if len(row.cells) > 1:
            right_cell = row.cells[1]
            new_right = right_text.strip()
            _replace_cell_text_keep_style(right_cell, new_right)


def apply_skills_paragraph_model(doc: Document, skills_model: Dict[str, Any]) -> None:
    """
    Rewrite SKILLS when it's a paragraph block after SKILLS/TECHNICAL SKILLS.
    """
    rows: List[List[str]] = skills_model.get("rows") or []
    if not rows:
        return

    paragraphs = list(doc.paragraphs)
    skills_heading_idx = find_heading_index(paragraphs, ["SKILLS", "TECHNICAL SKILLS"])
    if skills_heading_idx is None:
        return

    experience_idx = find_heading_index(
        paragraphs,
        ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
    )
    education_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])
    cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
    projects_idx = find_heading_index(
        paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "PROJECT", "PROJECT DETAILS"]
    )

    end_candidates = []
    for idx in [experience_idx, education_idx, cert_idx, projects_idx]:
        if idx is not None and idx > skills_heading_idx:
            end_candidates.append(idx)
    end_idx = min(end_candidates) if end_candidates else len(paragraphs)

    skill_paras_all = paragraphs[skills_heading_idx + 1: end_idx]
    skill_paras: List[DocxParagraph] = [p for p in skill_paras_all if p.text.strip()]
    if not skill_paras:
        return

    n_old = len(skill_paras)
    n_new = len(rows)
    n_common = min(n_old, n_new)

    def write_skill_line(para: DocxParagraph, left: str, right: str) -> None:
        template_run = para.runs[0] if para.runs else None

        for r in list(para.runs):
            r._element.getparent().remove(r._element)

        left = (left or "").strip()
        right = (right or "").strip()

        if left:
            r_left = para.add_run(left)
            if template_run is not None:
                copy_run_format(template_run, r_left)
            r_left.bold = True
        if left and right:
            r_colon = para.add_run(": ")
            if template_run is not None:
                copy_run_format(template_run, r_colon)
            r_colon.bold = False
        if right:
            r_right = para.add_run(right)
            if template_run is not None:
                copy_run_format(template_run, r_right)
            r_right.bold = False

    for i in range(n_common):
        left = rows[i][0] if len(rows[i]) > 0 else ""
        right = rows[i][1] if len(rows[i]) > 1 else ""
        write_skill_line(skill_paras[i], left, right)

    if n_new < n_old:
        for p in reversed(skill_paras[n_new:]):
            p._element.getparent().remove(p._element)

    if n_new > n_old:
        extra_rows = rows[n_old:]
        last_para = skill_paras[-1]
        for row in extra_rows:
            left = row[0] if len(row) > 0 else ""
            right = row[1] if len(row) > 1 else ""
            new_elem = deepcopy(last_para._element)
            last_para._element.addnext(new_elem)
            new_para = DocxParagraph(new_elem, last_para._parent)
            write_skill_line(new_para, left, right)
            last_para = new_para


def _write_label_value_paragraph(para: DocxParagraph, label: str, value: str) -> None:
    """
    For Environment / Role / Project Description lines:
      - label is bold
      - value is normal text after a space.
    """
    template_run = para.runs[0] if para.runs else None

    # Clear existing runs
    for r in list(para.runs):
        r._element.getparent().remove(r._element)

    label = (label or "").strip()
    value = (value or "").strip()

    # LABEL
    if label:
        r_label = para.add_run(label)
        if template_run is not None:
            copy_run_format(template_run, r_label)
        r_label.bold = True

    # SPACE
    if label and value:
        r_space = para.add_run(" ")
        if template_run is not None:
            copy_run_format(template_run, r_space)
        r_space.bold = False
        r_space.underline = False

    # VALUE
    if value:
        r_val = para.add_run(value)
        if template_run is not None:
            copy_run_format(template_run, r_val)
        r_val.bold = False
        r_val.underline = False


def apply_experience_model(doc: Document, experience_blocks: List[Dict[str, Any]]) -> None:
    """
    Rewrite the EXPERIENCE section according to the model.
    """
    if not experience_blocks:
        return

    paragraphs = list(doc.paragraphs)

    experience_idx = find_heading_index(
        paragraphs,
        ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
    )
    if experience_idx is None:
        return

    education_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])
    cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
    projects_idx = find_heading_index(
        paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "PROJECT", "PROJECT DETAILS"]
    )

    end_candidates = []
    for idx in [education_idx, cert_idx, projects_idx]:
        if idx is not None and idx > experience_idx:
            end_candidates.append(idx)
    exp_end_idx = min(end_candidates) if end_candidates else len(paragraphs)

    # detect job header indices
    job_header_indices: List[int] = []
    for idx in range(experience_idx + 1, exp_end_idx):
        if is_probable_job_header(paragraphs[idx]):
            job_header_indices.append(idx)

    if not job_header_indices:
        return

    num_jobs_to_update = min(len(job_header_indices), len(experience_blocks))

    # Process from last to first so deletions don't affect earlier indices
    for j in reversed(range(num_jobs_to_update)):
        paragraphs = list(doc.paragraphs)
        header_idx = job_header_indices[j]
        if header_idx >= len(paragraphs):
            continue

        header_para = paragraphs[header_idx]
        block = experience_blocks[j]

        # ---- Header ----
        new_header_text = (block.get("header") or header_para.text).strip()
        if new_header_text:
            replace_paragraph_text_keep_style(header_para, new_header_text)

        # Determine job range
        if j + 1 < len(job_header_indices):
            next_header_idx = job_header_indices[j + 1]
            job_end_idx = min(next_header_idx, exp_end_idx)
        else:
            job_end_idx = exp_end_idx

        bullet_paras: List[DocxParagraph] = []
        env_para: Optional[DocxParagraph] = None
        role_para: Optional[DocxParagraph] = None
        proj_para: Optional[DocxParagraph] = None

        for idx in range(header_idx + 1, job_end_idx):
            p = paragraphs[idx]
            t = p.text.strip()
            if not t:
                continue

            upper = t.upper()
            stripped_upper = upper.lstrip("•*-\u2022· ").strip()

            # Identify environment / role / project_description paragraphs
            if stripped_upper.startswith("ENVIRONMENT:"):
                if env_para is None:
                    env_para = p
                continue

            if stripped_upper.startswith("ROLE:"):
                if role_para is None:
                    role_para = p
                continue

            if stripped_upper.startswith("PROJECT DESCRIPTION:"):
                if proj_para is None:
                    proj_para = p
                continue

            # Case 2 (.NET): bold role line without 'Role:'
            if role_para is None and is_probable_role_line(p, t):
                role_para = p
                continue

            # Skip section-like headings
            if _is_heading_like(t):
                continue

            # Skip 'Responsibilities:' label – keep it as its own bold line, not a bullet
            if _is_responsibilities_label(t):
                continue

            bullet_paras.append(p)

        # ---- Update Environment ----
        env_model = block.get("environment")
        if env_para is not None and isinstance(env_model, dict):
            label = env_model.get("label", "") or "Environment:"
            value = env_model.get("value", "")
            _write_label_value_paragraph(env_para, label, value)

        # ---- Update / Insert Role ----
        role_model = block.get("role")
        if isinstance(role_model, dict):
            label = role_model.get("label", "") or "Role:"
            value = role_model.get("value", "")

            if role_para is None:
                # No explicit Role line in DOCX → create one after header
                new_p = deepcopy(header_para._p)
                header_para._p.addnext(new_p)
                role_para = DocxParagraph(new_p, header_para._parent)

            _write_label_value_paragraph(role_para, label, value)

        # ---- Update Project Description ----
        proj_model = block.get("project_description")
        if proj_para is not None and isinstance(proj_model, dict):
            label = proj_model.get("label", "") or "Project Description:"
            value = proj_model.get("value", "")
            _write_label_value_paragraph(proj_para, label, value)

        # ---- Update bullets ----
        raw_bullets: List[str] = block.get("bullets") or []

        # Do not send 'Responsibilities:' labels into bullet rewriting.
        new_bullets: List[str] = []
        for b in raw_bullets:
            if not b:
                continue
            if _is_responsibilities_label(b):
                # The actual Responsibilities paragraph stays as-is in the DOCX.
                continue
            new_bullets.append(b)

        if not bullet_paras and not new_bullets:
            continue

        if not new_bullets:
            for p in bullet_paras:
                elem = p._element
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
            continue

        update_bullet_paragraphs(bullet_paras, new_bullets)


def apply_header_model(doc: Document, header_model: Dict[str, Any], header_meta: Dict[str, Any]) -> None:
    """
    Apply header text changes while keeping styles (fonts/colors) as-is.

    We do a simple index-based replacement:
      - Do NOT change number of header lines.
      - Do NOT move them; only replace text.
    """
    lines: List[str] = header_model.get("lines") or []
    indices: List[int] = header_meta.get("indices") or []
    if not lines or not indices:
        return

    paragraphs = list(doc.paragraphs)
    n = min(len(lines), len(indices))
    for i in range(n):
        idx = indices[i]
        if idx < 0 or idx >= len(paragraphs):
            continue
        para = paragraphs[idx]
        replace_paragraph_text_keep_style(para, lines[i])


def apply_simple_lines_section(doc: Document, section_model: Dict[str, Any], section_meta: Dict[str, Any]) -> None:
    """
    For sections modeled as simple 'lines' + 'indices':
      - education
      - certifications
      - projects

    We treat those lines like bullets/paragraphs and use update_bullet_paragraphs
    so line counts can grow/shrink while preserving style.
    """
    lines: List[str] = section_model.get("lines") or []
    indices: List[int] = section_meta.get("indices") or []
    if not indices:
        return

    paragraphs = list(doc.paragraphs)
    bullet_paras: List[DocxParagraph] = [paragraphs[i] for i in indices if 0 <= i < len(paragraphs)]

    if not bullet_paras and not lines:
        return

    update_bullet_paragraphs(bullet_paras, lines)


def apply_other_sections_model(doc: Document, other_sections: List[Dict[str, Any]]) -> None:
    """
    Rewrite arbitrary 'other' sections detected by heading-like lines.

    Each section entry looks like:
      {
        "id": "awards",
        "title": "AWARDS",
        "kind": "lines",
        "lines": [...],
        "meta": {
          "heading_original_text": "AWARDS"
        }
      }

    We:
      - locate the heading paragraph by its original text (or current title),
      - treat all paragraphs until the next heading as this section's body,
      - *fully replace* those body paragraphs with new ones built from a style template,
      - keep paragraph styles intact but avoid reusing old runs (no leftover email/LinkedIn).
    """
    if not other_sections:
        return

    def build_heading_indices(paragraphs: List[DocxParagraph]) -> List[int]:
        """Heading indices in the *current* DOCX snapshot."""
        heading_idxs: List[int] = []
        for i, p in enumerate(paragraphs):
            txt = (p.text or "").strip()
            if not txt:
                continue

            # Skip obvious non-section labels
            if _is_responsibilities_label(txt):
                continue
            up = txt.upper()
            if up.startswith("ENVIRONMENT:") or up.startswith("ROLE:") or up.startswith("PROJECT DESCRIPTION:"):
                continue

            if _is_heading_like(txt):
                heading_idxs.append(i)

        heading_idxs.sort()
        return heading_idxs

    def _next_local_heading_after(idx: int, heading_indices: List[int]) -> Optional[int]:
        for hi in heading_indices:
            if hi > idx:
                return hi
        return None

    def _find_heading_index_for_section(
        paragraphs: List[DocxParagraph],
        original_title: str,
        new_title: str,
    ) -> Optional[int]:
        """
        Try to find the paragraph index whose text matches the original heading
        (first choice) or the new title (if user renamed it).
        """
        orig_norm = (original_title or "").strip().rstrip(" :").upper()
        new_norm = (new_title or "").strip().rstrip(" :").upper()

        # 1) Match by original heading text
        if orig_norm:
            for i, p in enumerate(paragraphs):
                txt = (p.text or "").strip().rstrip(" :").upper()
                if txt and txt == orig_norm:
                    return i

        # 2) Fallback: match by current title
        if new_norm:
            for i, p in enumerate(paragraphs):
                txt = (p.text or "").strip().rstrip(" :").upper()
                if txt and txt == new_norm:
                    return i

        return None

    for sec in other_sections:
        if not isinstance(sec, dict):
            continue

        kind = sec.get("kind") or "lines"
        if kind != "lines":
            # For now we only support simple line-based sections.
            continue

        title = (sec.get("title") or "").strip()
        meta = sec.get("meta") or {}
        if not isinstance(meta, dict):
            meta = {}

        original_title = (meta.get("heading_original_text") or "").strip()

        # Fresh snapshot of paragraphs for this section
        paragraphs = list(doc.paragraphs)
        if not paragraphs:
            return

        heading_idx = _find_heading_index_for_section(paragraphs, original_title, title)
        if heading_idx is None or heading_idx < 0 or heading_idx >= len(paragraphs):
            # Cannot safely locate this section in the DOCX – skip it.
            continue

        # Update the heading text itself (keep style)
        if title:
            heading_para = paragraphs[heading_idx]
            replace_paragraph_text_keep_style(heading_para, title)

        # Recompute heading indices on the *current* document
        paragraphs = list(doc.paragraphs)
        heading_indices = build_heading_indices(paragraphs)

        # Determine body range: everything between this heading and the next heading
        start = heading_idx + 1
        next_idx = _next_local_heading_after(heading_idx, heading_indices)
        end = next_idx if next_idx is not None else len(paragraphs)

        # Collect body paragraphs (non-empty) to rewrite
        paragraphs = list(doc.paragraphs)  # refresh again after heading change
        body_paras: List[DocxParagraph] = []
        for i in range(start, min(end, len(paragraphs))):
            p = paragraphs[i]
            if (p.text or "").strip():
                body_paras.append(p)

        lines = sec.get("lines") or []
        if not isinstance(lines, list):
            lines = []

        # If there is nothing to map and no paragraphs, skip
        if not body_paras and not lines:
            continue

        # If there are no existing body paragraphs, we don't know what style to clone.
        # In that case, just skip to avoid creating ugly default paragraphs.
        if not body_paras:
            continue

        # Make sure all body paragraphs share the same parent (same block/table cell).
        parents = {p._element.getparent() for p in body_paras if p._element is not None}
        if len(parents) != 1:
            # Fall back to generic updater if it's a weird mixed-parent case
            update_bullet_paragraphs(body_paras, [str(x) for x in lines])
            continue

        parent = parents.pop()
        # Use the first body paragraph as the style template
        template_para = body_paras[0]
        template_el = deepcopy(template_para._element)

        # Anchor: element before the first body paragraph
        first_elem = body_paras[0]._element
        anchor_before = first_elem.getprevious()

        # Remove all existing body paragraphs
        for p in body_paras:
            elem = p._element
            if elem is not None and elem.getparent() is parent:
                parent.remove(elem)

        # Helper to insert after a given element
        def insert_after(elem_to_insert, after_elem):
            if after_elem is None:
                parent.insert(0, elem_to_insert)
            else:
                parent.insert(parent.index(after_elem) + 1, elem_to_insert)

        last_inserted = anchor_before

        # Insert new body paragraphs, one per line, always wiping old runs
        clean_lines = [str(x).strip() for x in lines if str(x).strip()]
        if not clean_lines:
            # User effectively emptied this section's body: nothing to insert.
            continue

        for line in clean_lines:
            new_el = deepcopy(template_el)
            insert_after(new_el, last_inserted)
            last_inserted = new_el

            new_para = DocxParagraph(new_el, template_para._parent)
            # This removes all existing runs and builds fresh ones – no leftover email/LinkedIn.
            set_paragraph_text_with_bold_markers(new_para, line)




def apply_custom_sections_model(
    doc: Document,
    custom_sections: List[Dict[str, Any]],
    anchor_para: Optional[DocxParagraph] = None,
) -> None:
    """
    Append custom sections (heading + optional case-1 line + lines) at the end of the document.

    Each section dict is expected like:
      {
        "title": "Achievements",
        "lines": ["Did X", "Improved Y", ...]   # or "bullets": [...]
      }

    - Heading style is copied from an existing heading (SUMMARY / EXPERIENCE / etc.).
    - Content style is copied from a normal body/bullet paragraph (not contact lines).
    - For Case 1, if there is a 'line' paragraph under the first heading, we also
      clone that line under each custom section heading.
    """
    if not custom_sections:
        return

    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    # ---------- Find a template heading paragraph ----------
    heading_template: Optional[DocxParagraph] = None
    heading_index: Optional[int] = None
    for i, p in enumerate(paragraphs):
        txt = (p.text or "").strip()
        if not txt:
            continue
        if _is_heading_like(txt):
            heading_template = p
            heading_index = i
            break

    if heading_template is None or heading_index is None:
        return

    # ---------- Detect Case-1 style line right under the heading (optional) ----------
    line_template: Optional[DocxParagraph] = None
    if heading_index + 1 < len(paragraphs):
        cand = paragraphs[heading_index + 1]
        cand_text = (cand.text or "").strip()
        # Treat as a 'line' if it's empty or mostly dashes/underscores.
        if not cand_text or all(ch in "-_–—" for ch in cand_text):
            line_template = cand

    # ---------- Compute header_end to avoid contact lines for body template ----------
    summary_idx = find_heading_index(paragraphs, ["SUMMARY", "PROFESSIONAL SUMMARY"])
    skills_heading_idx = find_heading_index(paragraphs, ["SKILLS", "TECHNICAL SKILLS"])
    experience_idx = find_heading_index(
        paragraphs,
        ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
    )
    education_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])
    cert_idx = find_heading_index(paragraphs, ["CERTIFICATION", "CERTIFICATIONS"])
    project_idx = find_heading_index(paragraphs, ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"])

    first_section_candidates: List[int] = []
    for idx in [summary_idx, skills_heading_idx, experience_idx, education_idx, cert_idx, project_idx]:
        if idx is not None:
            first_section_candidates.append(idx)
    header_end = min(first_section_candidates) if first_section_candidates else 0

    # ---------- Find a template body paragraph (real bullet/content) ----------
    body_template: Optional[DocxParagraph] = None
    for idx in range(header_end, len(paragraphs)):
        p = paragraphs[idx]
        txt = (p.text or "").strip()
        if not txt:
            continue
        # Skip section headings
        if _is_heading_like(txt):
            continue
        lower = txt.lower()
        # Skip contact-style lines (email, LinkedIn, URLs)
        if "@" in lower or "linkedin" in lower or "http" in lower or "www." in lower:
            continue
        # Skip line-only paragraphs (all dashes/underscores)
        if all(ch in "-_–—" for ch in txt):
            continue
        body_template = p
        break

    # If we can't safely infer styles, bail out
    if body_template is None:
        return

    # Decide where to start inserting:
    # - If anchor_para is given -> insert BEFORE that paragraph.
    # - Otherwise -> append AFTER the last paragraph (original behavior).
    if paragraphs:
        if anchor_para is not None:
            base_para = anchor_para
            insert_before = True
        else:
            base_para = paragraphs[-1]
            insert_before = False
    else:
        return

    parent = base_para._element.getparent()
    if parent is None:
        return

    # Compute starting position in the parent XML children list
    insert_pos = parent.index(base_para._element)
    if not insert_before:
        insert_pos += 1

    def insert_next(elem_to_insert):
        nonlocal insert_pos
        parent.insert(insert_pos, elem_to_insert)
        insert_pos += 1

    for sec in custom_sections:
        if not isinstance(sec, dict):
            continue

        # Be flexible with field names coming from the UI
        title = (sec.get("title") or sec.get("name") or sec.get("heading") or "").strip()
        lines = sec.get("lines") or sec.get("bullets") or []
        lines = [str(x).strip() for x in lines if str(x).strip()]

        # Skip empty sections
        if not title and not lines:
            continue

        # ---- Heading (inherits style from an existing heading) ----
        new_heading_el = deepcopy(heading_template._element)
        insert_next(new_heading_el)
        new_heading_para = DocxParagraph(new_heading_el, heading_template._parent)
        replace_paragraph_text_keep_style(new_heading_para, title)

        # ---- Optional Case-1 line directly under the heading ----
        if line_template is not None:
            new_line_el = deepcopy(line_template._element)
            insert_next(new_line_el)
            # no text change – the border/line styling is in the paragraph formatting

        # ---- Lines under the heading (inherit style from a real body paragraph) ----
        for line in lines:
            new_body_el = deepcopy(body_template._element)
            insert_next(new_body_el)
            new_body_para = DocxParagraph(new_body_el, body_template._parent)
            # Allow **bold** markers in custom lines
            set_paragraph_text_with_bold_markers(new_body_para, line)


def _sanitize_xml_text(value: str) -> str:
    """
    Remove control characters that are invalid in XML/docx.
    """
    if not isinstance(value, str):
        return value

    out_chars = []
    for ch in value:
        code = ord(ch)
        # Allow tab, newline, carriage-return, and anything >= 0x20
        if code < 0x20 and ch not in ("\t", "\n", "\r"):
            continue
        out_chars.append(ch)

    return "".join(out_chars)


def _sanitize_model_for_docx(obj: Any) -> Any:
    """
    Recursively sanitize all strings in the model so they are safe for XML/docx.
    """
    if isinstance(obj, dict):
        return {k: _sanitize_model_for_docx(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize_model_for_docx(v) for v in obj]
    if isinstance(obj, str):
        return _sanitize_xml_text(obj)
    return obj


def _remove_section_by_heading(doc: Document, heading_keywords: List[str]) -> None:
    """
    Remove an entire section from the DOCX:
      - the heading paragraph that matches any of heading_keywords
      - all content paragraphs until the next section heading.

    Safe if the heading doesn't exist: it just does nothing.
    """
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    # Where does THIS section start?
    start_idx = find_heading_index(paragraphs, heading_keywords)
    if start_idx is None:
        return

    # All known section headings, reused for finding the "next" boundary
    all_heading_keyword_sets = [
        ["SUMMARY", "PROFESSIONAL SUMMARY"],
        ["SKILLS", "TECHNICAL SKILLS"],
        ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
        ["EDUCATION", "EDUCATION DETAILS"],
        ["CERTIFICATION", "CERTIFICATIONS"],
        ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"],
    ]

    end_idx = len(paragraphs)
    for kws in all_heading_keyword_sets:
        idx = find_heading_index(paragraphs, kws)
        if idx is not None and idx > start_idx and idx < end_idx:
            end_idx = idx

    # Remove paragraphs from bottom → top to keep indices valid
    for idx in range(end_idx - 1, start_idx - 1, -1):
        p = paragraphs[idx]
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)


def _remove_skills_section(doc: Document, skills_meta: Optional[Dict[str, Any]] = None) -> None:
    """
    Remove the SKILLS section.

    - If there's a SKILLS/TECHNICAL SKILLS heading, remove that heading
      and its content block (like other sections), AND also remove the
      skills table if present.

    - Otherwise, for table-only resumes, remove the table we originally
      treated as the skills table. If we don't know which one, fallback
      to the FIRST table.
    """
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    heading_keywords = ["SKILLS", "TECHNICAL SKILLS"]
    skills_idx = find_heading_index(paragraphs, heading_keywords)

    # 1) Heading-based skills section
    if skills_idx is not None:
        # Remove paragraphs under the SKILLS heading.
        _remove_section_by_heading(doc, heading_keywords)

        # Also remove the underlying skills table, if any.
        if not doc.tables:
            return

        table_index = None
        if skills_meta:
            try:
                table_index = skills_meta.get("table_index")
            except Exception:
                table_index = None

        if isinstance(table_index, int) and 0 <= table_index < len(doc.tables):
            table = doc.tables[table_index]
        else:
            # If table index doesn't match anymore, fallback to the first table.
            table = doc.tables[0]

        tbl = table._tbl
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)

        return

    # 2) Fallback: table-only skills (no explicit SKILLS heading)
    if not doc.tables:
        return

    table_index = None
    if skills_meta:
        try:
            table_index = skills_meta.get("table_index")
        except Exception:
            table_index = None

    if isinstance(table_index, int) and 0 <= table_index < len(doc.tables):
        table = doc.tables[table_index]
    else:
        table = doc.tables[0]

    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)


def _reorder_body_sections_by_order(doc: Document, section_order: List[str]) -> None:
    """
    Reorder the main body sections (Summary/Skills/Experience/Education/Certifications/Projects)
    according to the given section_order list from the model.

    Header stays where it is (everything before the first section heading).
    Custom sections are left wherever they are appended (usually at the end).
    """
    # Map our logical section IDs to heading keyword sets used in the document.
    # These MUST match the same heading lists used in parse_docx_to_model().
    section_headings_map: Dict[str, List[str]] = {
        "summary": ["SUMMARY", "PROFESSIONAL SUMMARY"],
        "skills": ["SKILLS", "TECHNICAL SKILLS"],
        "experience": [
            "PROFESSIONAL EXPERIENCE",
            "EXPERIENCE",
            "WORK EXPERIENCE",
            "EMPLOYMENT HISTORY",
        ],
        "education": ["EDUCATION", "EDUCATION DETAILS"],
        "certifications": ["CERTIFICATION", "CERTIFICATIONS"],
        "projects": ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"],
    }

    # Only consider IDs we know how to reorder
    valid_ids = [sid for sid in section_order if sid in section_headings_map]
    if not valid_ids:
        return

    paragraphs = list(doc.paragraphs)
    body = doc.element.body
    body_children = list(body)

    # Helper: map paragraph index -> index in body_children
    para_to_body_idx: Dict[int, int] = {}
    for i, p in enumerate(paragraphs):
        try:
            elem = p._p
        except AttributeError:
            continue
        if elem in body_children:
            para_to_body_idx[i] = body_children.index(elem)

    # Find the "start index" (in body_children) for each section heading
    section_start_body_idx: Dict[str, int] = {}
    for sid, keywords in section_headings_map.items():
        para_idx = find_heading_index(paragraphs, keywords)
        if para_idx is None:
            continue
        if para_idx in para_to_body_idx:
            section_start_body_idx[sid] = para_to_body_idx[para_idx]

    if not section_start_body_idx:
        # No recognizable headings; nothing to reorder.
        return

    # Compute the body index ranges [start, end) for each section block
    # We do this by sorting all found section starts and using the next start as the end.
    all_starts_sorted = sorted(section_start_body_idx.items(), key=lambda kv: kv[1])
    # Build a mapping section_id -> (start_idx, end_idx)
    section_ranges: Dict[str, Tuple[int, int]] = {}

    for idx, (sid, start_idx) in enumerate(all_starts_sorted):
        if idx + 1 < len(all_starts_sorted):
            next_start_idx = all_starts_sorted[idx + 1][1]
        else:
            next_start_idx = len(body_children)
        section_ranges[sid] = (start_idx, next_start_idx)

    # Determine the first section start to preserve "header" content above it
    first_section_start = min(start_idx for (_, (start_idx, _)) in section_ranges.items())

    # Collect header / top content (everything before the first section heading).
    header_prefix_children = body_children[:first_section_start]

    # Now assemble sections in the requested order, using the precomputed ranges.
    used_indices = set()
    ordered_section_children: List[object] = []

    for sid in valid_ids:
        rng = section_ranges.get(sid)
        if not rng:
            continue
        start_idx, end_idx = rng
        for i in range(start_idx, end_idx):
            if i in used_indices:
                continue
            ordered_section_children.append(body_children[i])
            used_indices.add(i)

    # Add any remaining body elements *after* all known sections.
    # This includes custom sections or any trailing content.
    trailing_children: List[object] = []
    for i, child in enumerate(body_children):
        if i < first_section_start:
            # Already accounted for in header_prefix_children
            continue
        if i in used_indices:
            continue
        trailing_children.append(child)

    new_body_children: List[object] = []
    new_body_children.extend(header_prefix_children)
    new_body_children.extend(ordered_section_children)
    new_body_children.extend(trailing_children)

    # Clear the body and re-append in new order
    for child in list(body):
        body.remove(child)
    for child in new_body_children:
        body.append(child)


def _remove_header_block(doc: Document) -> None:
    """
    Remove the top header block (name/title/contact) from the document.

    We treat everything before the first real section heading
    (SUMMARY / SKILLS / EXPERIENCE / EDUCATION / CERTS / PROJECTS)
    as 'header' and delete those elements (paragraphs + tables).
    """
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    # Find the first "real" section heading.
    heading_groups = [
        ["SUMMARY", "PROFESSIONAL SUMMARY"],
        ["SKILLS", "TECHNICAL SKILLS"],
        ["EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE"],
        ["EDUCATION"],
        ["CERTIFICATIONS", "CERTS"],
        ["PROJECTS"],
    ]

    first_heading_idx = None
    for idx, p in enumerate(paragraphs):
        text_upper = (p.text or "").strip().upper()
        if not text_upper:
            continue
        found = False
        for group in heading_groups:
            for kw in group:
                if kw in text_upper:
                    first_heading_idx = idx
                    found = True
                    break
            if found:
                break
        if found:
            break

    # If no heading is found, don't nuke the entire doc.
    if first_heading_idx is None:
        return

    # This is the first "real" heading paragraph
    first_heading_p = paragraphs[first_heading_idx]._p

    # Remove all body elements (paragraphs, tables, etc.)
    # that appear BEFORE that heading paragraph.
    body = doc.element.body
    for child in list(body):
        if child is first_heading_p:
            break
        body.remove(child)


def apply_model_to_docx(doc: Document, model: Dict[str, Any]) -> Document:
    """
    Apply the structured model back onto the DOCX.

    IMPORTANT BEHAVIOR:

    - If a section (summary/skills/experience/education/certifications/projects)
      IS PRESENT in the model => we rewrite it (same as before).
    - If that section is MISSING from the model => we treat that as
      "user deleted this section in the UI" and physically remove it
      from the DOCX (heading + content).

    This also works safely when the original resume never had that section:
    in that case there is no heading/table, so the remove helper is a no-op.
    """
    # ---------- HEADER: remove or sync based on model ----------
    paragraphs = list(doc.paragraphs)

    # CASE A: user deleted Header in UI  → no "header" key in model
    if "header" not in model:
        _remove_header_block(doc)

    else:
        # CASE B: header exists in model → either empty (treat as removed)
        #         or normal (sync role/title)
        header = model.get("header") or {}
        header_meta = model.get("header_meta") or {}

        if isinstance(header, dict):
            lines = header.get("lines") or []
            indices = header_meta.get("indices") or []

            # Normalize lines
            norm_lines = [(ln or "") for ln in lines]
            has_any_text = any((ln or "").strip() for ln in norm_lines)

            # B1) No real text or no indices → treat as removed
            if not has_any_text or not indices:
                _remove_header_block(doc)

            else:
                # Refresh paragraphs snapshot
                paragraphs = list(doc.paragraphs)

                # -------- CASE 2 / CASE 3: multiple header paragraphs (name + separate title) --------
                #   lines    = ["Name", "Role/Title", "Contact..."]
                #   indices  = [idx_name, idx_role, idx_contact, ...]
                if len(norm_lines) >= 2 and len(indices) >= 2:
                    new_title = (norm_lines[1] or "").strip()
                    if new_title:
                        idx_title = indices[1]
                        if 0 <= idx_title < len(paragraphs):
                            para = paragraphs[idx_title]
                            replace_paragraph_text_keep_style(para, new_title)

                # -------- CASE 1: single paragraph header with line breaks --------
                #   indices  = [idx_single_header_paragraph]
                #   lines from UI may be:
                #       ["Name", "Role/Title", "Contact..."]   (new behavior)
                #   OR legacy:
                #       ["Name\nRole\nContact..."]
                else:
                    idx_para = indices[0]
                    if 0 <= idx_para < len(paragraphs):
                        para = paragraphs[idx_para]

                        desired_role: Optional[str] = None

                        # Prefer explicit second line from UI if present
                        if len(norm_lines) >= 2 and (norm_lines[1] or "").strip():
                            desired_role = (norm_lines[1] or "").strip()
                        elif norm_lines:
                            # Legacy single-block style: split by newlines
                            block = norm_lines[0] or ""
                            parts = [p.strip() for p in block.split("\n")]
                            parts = [p for p in parts if p]
                            if len(parts) >= 2:
                                desired_role = parts[1]

                        if desired_role:
                            _update_single_paragraph_header_role(para, desired_role)

    # ---------- SUMMARY ----------
    if "summary" in model:
        summary_meta = model.get("summary_meta", {})
        apply_summary_model(doc, model["summary"], summary_meta)
    else:
        # User removed SUMMARY in UI (or resume never had it).
        _remove_section_by_heading(doc, ["SUMMARY", "PROFESSIONAL SUMMARY"])

    # ---------- SKILLS ----------
    # Always fetch skills_meta once so it's defined on all code paths.
    skills_meta = model.get("skills_meta")

    if "skills" in model:
        skills_model = model["skills"]

        # Detect "effectively empty" skills (all cells blank)
        raw_rows = skills_model.get("rows") or []
        cleaned_rows = [
            [(cell or "").strip() for cell in (row or [])]
            for row in raw_rows
        ]
        all_empty = (
            not cleaned_rows
            or all(not any(cell for cell in row) for row in cleaned_rows)
        )

        if all_empty:
            # Treat as "no skills section" → remove table/section.
            _remove_skills_section(doc, skills_meta)
        else:
            # Normal behavior: rewrite skills using the detected format.
            fmt = None
            if isinstance(skills_meta, dict):
                fmt = skills_meta.get("format")

            if not fmt:
                fmt = "table" if doc.tables else "paragraph"

            if fmt == "table":
                apply_skills_table_model(doc, skills_model)
            elif fmt == "paragraph":
                apply_skills_paragraph_model(doc, skills_model)
    else:
        # Only remove SKILLS if the parser originally detected a skills section.
        # If skills_meta is missing/empty, we assume the resume never had skills
        # and we should NOT delete any tables.
        if skills_meta:
            _remove_skills_section(doc, skills_meta)

    # Re-read paragraphs after possible removals
    paragraphs = list(doc.paragraphs)

    # ---------- EXPERIENCE ----------
    if "experience" in model:
        apply_experience_model(doc, model["experience"])
    else:
        _remove_section_by_heading(
            doc,
            ["PROFESSIONAL EXPERIENCE", "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
        )

    # ---------- EDUCATION ----------
    if "education" in model:
        edu_model = model.get("education") or {}

        # Check if the original DOCX has an EDUCATION heading
        paragraphs = list(doc.paragraphs)
        edu_idx = find_heading_index(paragraphs, ["EDUCATION", "EDUCATION DETAILS"])

        if edu_idx is None:
            # No explicit EDUCATION section in the DOCX.
            # This is the case for your Case 1 template where the degree
            # lives under a "DETAILS" heading at the bottom.
            edu_lines = edu_model.get("lines") or []

            if edu_lines:
                other_sections = model.get("other_sections") or []
                for sec in other_sections:
                    if not isinstance(sec, dict):
                        continue
                    title = (sec.get("title") or "").strip().upper()
                    # Treat "DETAILS" as the host for education lines
                    if title == "DETAILS":
                        # Overwrite DETAILS body with the education lines
                        sec["lines"] = list(edu_lines)
                        break

        # Still call the normal applier – it will safely no-op
        # if there really is no EDUCATION heading.
        apply_education_model(doc, edu_model)
    else:
        _remove_section_by_heading(doc, ["EDUCATION", "EDUCATION DETAILS"])

    # ---------- CERTIFICATIONS ----------
    if "certifications" in model:
        apply_certifications_model(doc, model["certifications"])
    else:
        _remove_section_by_heading(doc, ["CERTIFICATION", "CERTIFICATIONS"])

    # ---------- PROJECTS ----------
    if "projects" in model:
        apply_projects_model(doc, model["projects"])
    else:
        _remove_section_by_heading(
            doc,
            ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"],
        )

    # ---------- OTHER / GENERIC SECTIONS ----------
    other_sections = model.get("other_sections") or []
    if other_sections:
        apply_other_sections_model(doc, other_sections)

    # ---------- SECTION ORDER (reorder core blocks for Summary/Skills/Experience/...) ----------
    section_order = model.get("section_order")
    if isinstance(section_order, list) and section_order:
        _reorder_body_sections_by_order(doc, section_order)

    # ---------- CUSTOM SECTIONS (insert according to section_order) ----------
    custom_sections = model.get("custom_sections") or model.get("extra_sections") or []
    if custom_sections:
        anchor_para = None

        if isinstance(section_order, list) and "custom_sections" in section_order:
            # known core section ids
            recognized_ids = [
                "summary",
                "skills",
                "experience",
                "education",
                "certifications",
                "projects",
            ]

            # Heading keywords (same as in reorder + parser)
            headings_map: Dict[str, List[str]] = {
                "summary": ["SUMMARY", "PROFESSIONAL SUMMARY"],
                "skills": ["SKILLS", "TECHNICAL SKILLS"],
                "experience": [
                    "PROFESSIONAL EXPERIENCE",
                    "EXPERIENCE",
                    "WORK EXPERIENCE",
                    "EMPLOYMENT HISTORY",
                ],
                "education": ["EDUCATION", "EDUCATION DETAILS"],
                "certifications": ["CERTIFICATION", "CERTIFICATIONS"],
                "projects": ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS"],
            }

            try:
                idx_custom = section_order.index("custom_sections")
            except ValueError:
                idx_custom = -1

            if idx_custom != -1:
                next_core_id = None
                for sid in section_order[idx_custom + 1:]:
                    if sid in recognized_ids:
                        next_core_id = sid
                        break

                if next_core_id is not None:
                    paragraphs = list(doc.paragraphs)
                    para_idx = find_heading_index(
                        paragraphs, headings_map.get(next_core_id, [])
                    )
                    if para_idx is not None and 0 <= para_idx < len(paragraphs):
                        anchor_para = paragraphs[para_idx]

        # If anchor_para is None → append at end (default behavior)
        apply_custom_sections_model(doc, custom_sections, anchor_para=anchor_para)

    return doc


def apply_model_and_generate_docx(doc_bytes: bytes, model: Dict[str, Any]) -> bytes:
    """
    Takes original DOCX bytes + edited model and returns updated DOCX bytes.

    Before applying, we sanitize all strings in the model to remove XML-invalid
    control characters.
    """
    doc = load_document_from_bytes(doc_bytes)

    # Clean any bad control characters from the model text
    clean_model = _sanitize_model_for_docx(model)

    # Apply the cleaned model to the doc
    updated_doc = apply_model_to_docx(doc, clean_model)

    out_buf = BytesIO()
    updated_doc.save(out_buf)
    return out_buf.getvalue()


def build_cover_letter_docx(model: Dict[str, Any], cover_letter_text: str) -> bytes:
    """
    Build a standalone DOCX cover letter with a simple, standard template:

      - 1 inch margins on all sides
      - Header using the resume header lines (name + contact info)
      - Body paragraphs from the generated cover letter text

    We assume `cover_letter_text` already contains the greeting and sign-off
    (e.g., "Dear Hiring Manager," ... "Sincerely, <Name>").
    """
    doc = Document()

    # Standard 1-inch margins
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Base style: Calibri 11, modest spacing
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)    # default gap
        pf.line_spacing = 1.15
    except Exception:
        pass  # fall back to Word defaults if anything explodes

    # ---- Header: name + contact info ----
    header_lines: List[str] = []
    if isinstance(model, dict):
        header = model.get("header") or {}
        if isinstance(header, dict):
            lines = header.get("lines") or []
            if isinstance(lines, list):
                for x in lines:
                    s = str(x).strip()
                    if s:
                        header_lines.append(s)

    if header_lines:
        name = header_lines[0]

        name_p = doc.add_paragraph()
        name_run = name_p.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_p.paragraph_format.space_before = Pt(0)
        name_p.paragraph_format.space_after = Pt(0)

        if len(header_lines) > 1:
            for line in header_lines[1:]:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

        # One controlled blank gap after the header block
        gap_p = doc.add_paragraph()
        gap_p.paragraph_format.space_before = Pt(6)
        gap_p.paragraph_format.space_after = Pt(6)

    # ---- Body: each non-empty line becomes its own paragraph ----
    text = (cover_letter_text or "").strip()
    if text:
        lines = [ln.strip() for ln in text.splitlines()]
        # Treat these as closings (case-insensitive)
        closing_prefixes = ("sincerely", "best regards", "regards")

        i = 0
        n = len(lines)
        while i < n:
            line = lines[i]
            if not line:
                i += 1
                continue

            low = line.lower()

            # Closing block: "Sincerely," / "Regards," + name on next line
            if any(low.startswith(pref) for pref in closing_prefixes):
                # Extra space before closing
                extra = doc.add_paragraph()
                extra.paragraph_format.space_before = Pt(12)
                extra.paragraph_format.space_after = Pt(0)

                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)  # keep closing tight

                # First line: "Sincerely," or "Regards,"
                run = p.add_run(line)

                # Find next non-empty line as the name, if present
                j = i + 1
                while j < n and not lines[j].strip():
                    j += 1
                if j < n:
                    name_line = lines[j].strip()
                    p.add_run().add_break()   # Shift+Enter style
                    p.add_run(name_line)
                    i = j  # skip name line in outer loop

            else:
                # Normal paragraph: one line = one paragraph with some gap after
                p = doc.add_paragraph(line)
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(8)  # visible gap after each paragraph

            i += 1

    out_buf = BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue()
